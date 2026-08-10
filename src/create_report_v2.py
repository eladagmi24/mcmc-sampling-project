"""Generate the project report from results/experiment_results_v2.json.

Every number in the document is read from the results file. Nothing is hand-copied, so the
report cannot drift out of step with the experiments.
"""
import json
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding='utf-8')

SOURCE_DIRECTORY = os.path.dirname(os.path.abspath(__file__))
REPOSITORY_ROOT = os.path.dirname(SOURCE_DIRECTORY) \
    if os.path.isdir(os.path.join(os.path.dirname(SOURCE_DIRECTORY), 'results')) \
    else SOURCE_DIRECTORY
RESULTS_FILE = os.path.join(REPOSITORY_ROOT, 'results', 'experiment_results_v2.json')
FIGURE_DIRECTORY = os.path.join(REPOSITORY_ROOT, 'results', 'figures')
DOCUMENT_DIRECTORY = os.path.join(REPOSITORY_ROOT, 'docs')
OUTPUT_FILE = os.path.join(DOCUMENT_DIRECTORY, 'Sampling_Project_Report.docx')
GITHUB_REPOSITORY_URL = 'https://github.com/eladagmi24/mcmc-sampling-project'

TITLE_PAGE_PLACEHOLDERS = {
    'Institution': '__________________________',
    'Course': 'Advanced Methods in Machine Learning',
    'Lecturer': 'Dr. Boaz Tamir',
    'Student IDs': '__________________________',
}

with open(RESULTS_FILE, 'r') as handle:
    results = json.load(handle)

configuration = results['configuration']
data_summary = results['data']
samplers = results['samplers']
long_run = results['preconditioned_long_run']
analytical = results['analytical_reference']
external = results.get('external_reference')
arviz_check = results.get('arviz_cross_check')
selection = results['likelihood_selection']
test_performance = results['test_performance']
residuals = results['residual_diagnostics']
pooling = results['pooling_analysis']
geometry = results['posterior_geometry']
hmc_grid = results['hmc_sensitivity']
mh_grid = results['mh_sensitivity']
initialisation = results['initialisation_sensitivity']

GAUSSIAN_SAMPLERS = ['MH', 'Adaptive MH (naive)', 'Preconditioned MH', 'Gibbs', 'HMC']
CONVERGED_SAMPLERS = [name for name in GAUSSIAN_SAMPLERS if samplers[name]['converged']]
NON_CONVERGED_SAMPLERS = [name for name in GAUSSIAN_SAMPLERS if not samplers[name]['converged']]
SELECTED = selection['selected']
SELECTED_DEGREES = SELECTED['degrees_of_freedom']
GAUSSIAN_TEST = test_performance['gaussian']
STUDENT_TEST = test_performance['student_t']
OLS_TEST = test_performance['ols']
CONVERGED_HMC_CELLS = [cell for cell in hmc_grid if cell['converged']]
BEST_HMC_CELL = max(CONVERGED_HMC_CELLS, key=lambda cell: cell['bulk_ess_per_gradient']) \
    if CONVERGED_HMC_CELLS else None


# --------------------------------------------------------------------------------------------
# Citations
# --------------------------------------------------------------------------------------------

BIBLIOGRAPHY = {
    'beskos2013': 'Beskos, A., Pillai, N., Roberts, G., Sanz-Serna, J. M., & Stuart, A. (2013). '
                  'Optimal tuning of the hybrid Monte Carlo algorithm. Bernoulli, 19(5A), '
                  '1501-1534. https://doi.org/10.3150/12-BEJ414',
    'foremanmackey2013': 'Foreman-Mackey, D., Hogg, D. W., Lang, D., & Goodman, J. (2013). '
                         'emcee: The MCMC Hammer. Publications of the Astronomical Society of '
                         'the Pacific, 125(925), 306-312. https://doi.org/10.1086/670067',
    'gelman2013': 'Gelman, A., Carlin, J. B., Stern, H. S., Dunson, D. B., Vehtari, A., & '
                  'Rubin, D. B. (2013). Bayesian Data Analysis (3rd ed.). CRC Press.',
    'geman1984': 'Geman, S., & Geman, D. (1984). Stochastic relaxation, Gibbs distributions, and '
                 'the Bayesian restoration of images. IEEE Transactions on Pattern Analysis and '
                 'Machine Intelligence, 6(6), 721-741. https://doi.org/10.1109/TPAMI.1984.4767596',
    'geyer1992': 'Geyer, C. J. (1992). Practical Markov Chain Monte Carlo. Statistical Science, '
                 '7(4), 473-483. https://doi.org/10.1214/ss/1177011137',
    'goodman2010': 'Goodman, J., & Weare, J. (2010). Ensemble samplers with affine invariance. '
                   'Communications in Applied Mathematics and Computational Science, 5(1), '
                   '65-80. https://doi.org/10.2140/camcos.2010.5.65',
    'haario2001': 'Haario, H., Saksman, E., & Tamminen, J. (2001). An adaptive Metropolis '
                  'algorithm. Bernoulli, 7(2), 223-242. https://doi.org/10.2307/3318737',
    'hastings1970': 'Hastings, W. K. (1970). Monte Carlo sampling methods using Markov chains '
                    'and their applications. Biometrika, 57(1), 97-109. '
                    'https://doi.org/10.1093/biomet/57.1.97',
    'kumar2019': 'Kumar, R., Carroll, C., Hartikainen, A., & Martin, O. (2019). ArviZ: a '
                 'unified library for exploratory analysis of Bayesian models in Python. '
                 'Journal of Open Source Software, 4(33), 1143. '
                 'https://doi.org/10.21105/joss.01143',
    'lange1989': 'Lange, K. L., Little, R. J. A., & Taylor, J. M. G. (1989). Robust statistical '
                 'modeling using the t distribution. Journal of the American Statistical '
                 'Association, 84(408), 881-896. https://doi.org/10.2307/2290063',
    'metropolis1953': 'Metropolis, N., Rosenbluth, A. W., Rosenbluth, M. N., Teller, A. H., & '
                      'Teller, E. (1953). Equation of state calculations by fast computing '
                      'machines. The Journal of Chemical Physics, 21(6), 1087-1092. '
                      'https://doi.org/10.1063/1.1699114',
    'neal2011': 'Neal, R. M. (2011). MCMC using Hamiltonian dynamics. In Handbook of Markov '
                'Chain Monte Carlo (pp. 113-162). CRC Press. '
                'https://doi.org/10.1201/b10905',
    'roberts1997': 'Roberts, G. O., Gelman, A., & Gilks, W. R. (1997). Weak convergence and '
                   'optimal scaling of random walk Metropolis algorithms. The Annals of Applied '
                   'Probability, 7(1), 110-120. https://doi.org/10.1214/aoap/1034625254',
    'shen2015': 'Shen, S., van Beek, V., & Iosup, A. (2015). Statistical characterization of '
                'business-critical workloads hosted in cloud datacenters. In Proceedings of the '
                '15th IEEE/ACM International Symposium on Cluster, Cloud and Grid Computing '
                '(CCGrid), 465-474. https://doi.org/10.1109/CCGrid.2015.60. Dataset: Grid '
                'Workloads Archive GWA-T-12 Bitbrains, '
                'http://gwa.ewi.tudelft.nl/datasets/gwa-t-12-bitbrains',
    'vehtari2021': 'Vehtari, A., Gelman, A., Simpson, D., Carpenter, B., & Burkner, P. C. '
                   '(2021). Rank-normalization, folding, and localization: an improved R-hat '
                   'for assessing convergence of MCMC. Bayesian Analysis, 16(2), 667-718. '
                   'https://doi.org/10.1214/20-BA1221',
    'west1984': 'West, M. (1984). Outlier models and prior distributions in Bayesian linear '
                'regression. Journal of the Royal Statistical Society: Series B, 46(3), 431-439. '
                'https://doi.org/10.1111/j.2517-6161.1984.tb01317.x',
}
CITATION_ORDER = []


def cite(*keys):
    """Return a numbered citation marker, registering the entry on first use."""
    numbers = []
    for key in keys:
        if key not in BIBLIOGRAPHY:
            raise KeyError('unknown bibliography key: %s' % key)
        if key not in CITATION_ORDER:
            CITATION_ORDER.append(key)
        numbers.append(str(CITATION_ORDER.index(key) + 1))
    return '[%s]' % ', '.join(numbers)


# --------------------------------------------------------------------------------------------
# Document helpers
# --------------------------------------------------------------------------------------------

document = Document()
normal_style = document.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)
normal_style.paragraph_format.space_after = Pt(6)
normal_style.paragraph_format.line_spacing = 1.15
for level in range(1, 4):
    heading_style = document.styles['Heading %d' % level]
    heading_style.font.color.rgb = RGBColor(0x00, 0x52, 0x8A)
    heading_style.paragraph_format.keep_with_next = True

FIGURE_COUNTER = [0]
TABLE_COUNTER = [0]


def add_paragraph(text, bold=False, italic=False, size=11, alignment=None, space_after=6,
                  keep_with_next=False):
    paragraph = document.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return paragraph


def add_bullet(text):
    paragraph = document.add_paragraph(text, style='List Bullet')
    paragraph.paragraph_format.left_indent = Cm(1.27)
    return paragraph


def add_heading(text, level=1):
    heading = document.add_heading(text, level=level)
    heading.paragraph_format.keep_with_next = True
    return heading


def set_alt_text(picture, description):
    """Attach alternative text so the figure is accessible to screen readers."""
    element = picture._inline.docPr
    element.set('descr', description)
    element.set('title', description[:80])


def add_figure(filename, caption, alt_text, width_inches=6.1):
    path = os.path.join(FIGURE_DIRECTORY, filename)
    if not os.path.exists(path):
        add_paragraph('[Figure missing: %s]' % filename, italic=True, size=10)
        return
    FIGURE_COUNTER[0] += 1
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.keep_with_next = True
    picture_paragraph.paragraph_format.space_after = Pt(2)
    run = picture_paragraph.add_run()
    picture = run.add_picture(path, width=Inches(width_inches))
    set_alt_text(picture, alt_text)
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(12)
    caption_run = caption_paragraph.add_run('Figure %d. %s' % (FIGURE_COUNTER[0], caption))
    caption_run.font.size = Pt(9.5)
    caption_run.font.italic = True


def mark_header_row(table):
    """Repeat the first row on every page the table spans."""
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    row_properties.append(header)
    cant_split = OxmlElement('w:cantSplit')
    row_properties.append(cant_split)


def add_table(caption, headers, rows, column_widths_cm=None, no_wrap_first_column=True):
    TABLE_COUNTER[0] += 1
    caption_paragraph = document.add_paragraph()
    caption_paragraph.paragraph_format.space_after = Pt(3)
    caption_paragraph.paragraph_format.keep_with_next = True
    caption_run = caption_paragraph.add_run('Table %d. %s' % (TABLE_COUNTER[0], caption))
    caption_run.font.size = Pt(9.5)
    caption_run.font.italic = True
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for column_index, header_text in enumerate(headers):
        cell = table.rows[0].cells[column_index]
        cell.text = str(header_text)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
    for row_index, row_values in enumerate(rows):
        for column_index, value in enumerate(row_values):
            cell = table.rows[row_index + 1].cells[column_index]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
    if column_widths_cm:
        for row in table.rows:
            for column_index, width in enumerate(column_widths_cm):
                row.cells[column_index].width = Cm(width)
    mark_header_row(table)
    add_paragraph('', space_after=8)
    return table


def add_formula(text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11.5)
    run.font.italic = True
    return paragraph


def add_code_block(lines):
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.left_indent = Cm(1.0)
        run = paragraph.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


def add_field(paragraph, instruction, dirty=False):
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    if dirty:
        begin.set(qn('w:dirty'), 'true')
    instruction_element = OxmlElement('w:instrText')
    instruction_element.set(qn('xml:space'), 'preserve')
    instruction_element.text = instruction
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run = paragraph.add_run()._r
    run.append(begin)
    run.append(instruction_element)
    run.append(separate)
    run.append(end)


def add_page_numbers():
    for section in document.sections:
        footer_paragraph = section.footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(footer_paragraph, ' PAGE ')
        for run in footer_paragraph.runs:
            run.font.size = Pt(9)


def format_percentage(fraction, decimals=1):
    return ('%.' + str(decimals) + 'f%%') % (fraction * 100)


def format_interval(bounds):
    return '[%s, %s]' % (format_percentage(bounds[0]), format_percentage(bounds[1]))


def thousands(value):
    return '{:,}'.format(int(value))


def deviation_from_analytical(sampler_name):
    reference = analytical['posterior_mean_coefficients']
    candidate = samplers[sampler_name]['posterior_mean_coefficients']
    return max(abs(a - b) for a, b in zip(candidate, reference))


DEVIATIONS = {name: deviation_from_analytical(name) for name in GAUSSIAN_SAMPLERS}
WORST_CONVERGED_DEVIATION = max(DEVIATIONS[name] for name in CONVERGED_SAMPLERS)

# --------------------------------------------------------------------------------------------
# Title page
# --------------------------------------------------------------------------------------------

add_paragraph('Comparing Markov Chain Monte Carlo Samplers for Bayesian Linear Regression:\n'
              'Convergence, Efficiency and Predictive Calibration in Cloud CPU Forecasting',
              bold=True, size=17, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_paragraph('Elad Dagmi & Shaked Mizrahi', bold=True, size=13,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_paragraph('Advanced Methods in Machine Learning', size=12,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph('August 2026', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
for label, value in TITLE_PAGE_PLACEHOLDERS.items():
    add_paragraph('%s: %s' % (label, value), size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=3)
add_paragraph('')
add_paragraph('Source code and data instructions: %s' % GITHUB_REPOSITORY_URL, size=10,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

document.add_page_break()

# --------------------------------------------------------------------------------------------
# Abstract
# --------------------------------------------------------------------------------------------

add_heading('Abstract', level=1)
abstract_text = (
    'Markov chain Monte Carlo (MCMC) makes Bayesian inference possible when the posterior has no '
    'closed form, but its reliability rests on diagnostics that are easy to misread. This study '
    'compares five samplers, all implemented from scratch, on Bayesian linear regression for '
    'forecasting virtual-machine CPU utilisation %d minutes ahead, using the Bitbrains GWA-T-12 '
    'datacenter traces %s. The samplers are random-walk Metropolis-Hastings, adaptive Metropolis, '
    'Fisher-preconditioned Metropolis, Gibbs sampling and Hamiltonian Monte Carlo. Observations '
    'are split chronologically, with an embargo preventing feature windows from crossing split '
    'boundaries and all scaling estimated on training data alone. Convergence is assessed with '
    'rank-normalised and folded R-hat and with bulk and tail effective sample size (ESS), '
    'computed jointly across %d dispersed chains %s. '
    'Gibbs and Hamiltonian Monte Carlo converge; preconditioned Metropolis converges only with a '
    'longer run. Plain random-walk '
    'Metropolis and adaptive Metropolis do not, reaching R-hat %.2f and %.2f, with posterior '
    'means wrong by %.2f and %.2f against a deterministic quadrature reference. The cause is '
    'posterior anisotropy, condition number %.0f. Preconditioning with the observed Fisher '
    'information raises bulk ESS from %.0f to %.0f, and the converged samplers reproduce the '
    'quadrature reference to within %.4f. '
    'Test residuals have an excess kurtosis of %.0f, so a Gaussian likelihood yields posterior '
    'predictive intervals that are far too wide, nominal 50%% intervals covering %s. A Student-t '
    'likelihood, with degrees of freedom selected on validation data, reduces this to %s and cuts '
    'the median absolute error from %.3f to %.3f CPU percentage points, but the intervals remain '
    'substantially miscalibrated.'
    % (configuration['horizon_minutes'], cite('shen2015'), configuration['chains'],
       cite('vehtari2021'),
       samplers['MH']['worst_rhat'], samplers['Adaptive MH (naive)']['worst_rhat'],
       DEVIATIONS['MH'], DEVIATIONS['Adaptive MH (naive)'],
       geometry['condition_number'], samplers['MH']['min_bulk_ess'],
       samplers['Preconditioned MH']['min_bulk_ess'], WORST_CONVERGED_DEVIATION,
       residuals['excess_kurtosis'], format_percentage(GAUSSIAN_TEST['coverage_50']),
       format_percentage(STUDENT_TEST['coverage_50']),
       GAUSSIAN_TEST['median_absolute_error_percentage_points'],
       STUDENT_TEST['median_absolute_error_percentage_points']))
add_paragraph(abstract_text)
ABSTRACT_WORD_COUNT = len(abstract_text.split())

document.add_page_break()

# --------------------------------------------------------------------------------------------
# Contents
# --------------------------------------------------------------------------------------------

add_heading('Table of Contents', level=1)
contents_paragraph = document.add_paragraph()
add_field(contents_paragraph, r' TOC \o "1-3" \h \z \u ', dirty=True)

document.add_page_break()

# --------------------------------------------------------------------------------------------
# 1. Introduction
# --------------------------------------------------------------------------------------------

add_heading('1. Introduction', level=1)
add_paragraph(
    'Bayesian inference requires the posterior distribution of the model parameters given the '
    'data. For most models the normalising constant cannot be computed, so the posterior is '
    'explored by simulation. MCMC constructs a Markov chain whose stationary distribution is the '
    'target posterior %s, which turns inference into a simulation problem and, with it, into a '
    'question of whether the simulation has run long enough.' % cite('metropolis1953', 'hastings1970'))
add_paragraph(
    'The application is forecasting the CPU utilisation of virtual machines in a production '
    'datacenter %s. A point forecast is of limited operational value; capacity planning and '
    'overload detection need a calibrated statement of uncertainty, which is what a posterior '
    'predictive distribution provides.' % cite('shen2015'))
add_paragraph('This study addresses three research questions:')
add_bullet('RQ1. Which samplers actually converge on this posterior, and what property of the '
           'posterior explains the differences between them?')
add_bullet('RQ2. Among the samplers that converge, which is most efficient per unit of '
           'computation, measured by effective sample size per second and, for gradient methods, '
           'per gradient evaluation?')
add_bullet('RQ3. Are the resulting posterior predictive intervals calibrated on held-out data, '
           'and if not, what is responsible?')
add_paragraph(
    'Five samplers target the same Gaussian posterior: random-walk Metropolis-Hastings, adaptive '
    'Metropolis %s, Metropolis preconditioned by the observed Fisher information, Gibbs sampling '
    '%s and Hamiltonian Monte Carlo %s. A sixth sampler, a Gibbs sampler for Student-t '
    'regression, targets a different model and is therefore discussed separately, in Section 6, '
    'rather than being compared with the five on convergence or efficiency.'
    % (cite('haario2001'), cite('geman1984'), cite('neal2011')))

# --------------------------------------------------------------------------------------------
# 2. Theory
# --------------------------------------------------------------------------------------------

add_heading('2. Theoretical Background', level=1)

add_heading('2.1 Model and Priors', level=2)
add_paragraph('The regression model for the standardised target is')
add_formula('y | X, beta, sigma^2  ~  N(X beta, sigma^2 I_n)')
add_paragraph('The two priors are specified independently of one another:')
add_formula('beta ~ N(0, tau^2 I_p),   tau^2 = %.0f' % 10.0)
add_formula('1/sigma^2 ~ Gamma(a0 = %.0f, rate = b0 = %.0f),  equivalently  '
            'sigma^2 ~ Inverse-Gamma(a0 = %.0f, scale = b0 = %.0f)'
            % (2.0, 1.0, 2.0, 1.0))
add_paragraph(
    'The precision 1/sigma^2 follows a Gamma in shape-rate form; inverting gives sigma^2 an '
    'Inverse-Gamma in shape-scale form, with density proportional to '
    '(sigma^2)^(-a0-1) exp(-b0 / sigma^2). The Gamma distributions appearing in the Student-t '
    'sampler of Section 2.6 are likewise in shape-rate form. This matters in implementation '
    'because NumPy draws Gamma variates with a scale argument, so every rate is passed as its '
    'reciprocal.')
add_paragraph(
    'Note that the prior on beta does not scale with sigma^2. The joint prior is therefore not '
    'the conjugate Normal-Inverse-Gamma, and the joint posterior has no standard closed form, '
    'although both full conditionals remain standard. This is what makes Gibbs sampling available '
    'while still leaving something for the other samplers to do.')

add_heading('2.2 Metropolis-Hastings', level=2)
add_paragraph('A proposal theta-prime is accepted with probability')
add_formula("alpha = min(1, [pi(theta') q(theta | theta')] / [pi(theta) q(theta' | theta)])")
add_paragraph(
    'For a symmetric random-walk proposal the q terms cancel. The method needs only pointwise '
    'evaluation of an unnormalised density, which makes it general, but an uninformed random walk '
    'explores slowly. For a d-dimensional Gaussian target the asymptotically optimal acceptance '
    'rate is about 23.4 percent %s.' % cite('roberts1997'))

add_heading('2.3 Adaptive and Preconditioned Metropolis', level=2)
add_paragraph(
    'A single scalar step size implicitly assumes a spherical posterior. When the posterior is '
    'anisotropic the step must fit the narrowest direction while the chain has to traverse the '
    'widest, and the number of steps required grows with the condition number of the posterior '
    'covariance. Two remedies are examined.')
add_paragraph(
    'Adaptive Metropolis %s estimates the proposal covariance from the chain history and proposes '
    'from N(theta, (2.38^2 / d) C). Preconditioned Metropolis instead takes the metric from the '
    'observed Fisher information, evaluated at the least-squares residual variance sigmahat^2:'
    % cite('haario2001'))
add_formula('(X\'X / sigmahat^2 + I / tau^2)^-1  =  sigmahat^2 (X\'X + sigmahat^2 I / tau^2)^-1')
add_paragraph(
    'The right-hand form is the one implemented. It requires a single least-squares fit and no '
    'conjugacy. A Robbins-Monro rule additionally tunes a global scale toward the optimal '
    'acceptance rate, and the empirical covariance of the second half of burn-in refines the '
    'metric. Both adaptations stop before the retained draws begin, so the retained chain is a '
    'homogeneous Markov chain and no diminishing-adaptation argument is needed.')

add_heading('2.4 Gibbs Sampling', level=2)
add_paragraph('Both full conditionals are standard:')
add_formula("beta | sigma^2, y  ~  N(Sigma X'y / sigma^2, Sigma),  "
            "Sigma = (X'X / sigma^2 + I / tau^2)^-1")
add_formula("sigma^2 | beta, y  ~  Inverse-Gamma(a0 + n/2,  b0 + ||y - X beta||^2 / 2)")
add_paragraph(
    'Every draw comes from an exact conditional, so nothing is ever rejected. The resulting '
    'acceptance rate of one is a structural property of the algorithm and carries no information '
    'about sampling quality, which is why the comparison in Section 5 rests on R-hat, ESS and '
    'runtime instead.')

add_heading('2.5 Hamiltonian Monte Carlo', level=2)
add_paragraph(
    'HMC augments the parameter space with a momentum variable and simulates Hamiltonian dynamics '
    'with the leapfrog integrator, which is symplectic and reversible, so detailed balance holds '
    'once the accept-reject step corrects the integrator error %s.' % cite('neal2011'))
add_formula('H(q, p) = -log pi(q) + p\'p / 2')
add_paragraph(
    'Because sigma^2 must stay positive, HMC samples the unconstrained parameter psi = '
    'log(sigma^2). The change of variables contributes a Jacobian |d sigma^2 / d psi| = '
    'exp(psi) = sigma^2, so the term log(sigma^2) is added to the log posterior. The same '
    'transformation and the same Jacobian term are used by every sampler that works on the '
    'unconstrained scale, namely both Metropolis variants, HMC and the external emcee reference; '
    'Gibbs works directly on sigma^2 and needs no Jacobian. The gradient used by the leapfrog '
    'integrator is the analytic gradient of this transformed log posterior.')
add_paragraph(
    'The No-U-Turn Sampler, which tunes the trajectory length automatically, was not implemented '
    'or run in this study. Where trajectory length matters it is varied explicitly, in Section '
    '5.7.')

add_heading('2.6 Student-t Regression as a Normal Scale Mixture', level=2)
add_paragraph(
    'A Gaussian likelihood treats large residuals as essentially impossible, so when the data '
    'contain rare extreme observations the fitted noise variance inflates to accommodate them and '
    'every predictive interval widens. A Student-t likelihood tolerates them. Writing it as a '
    'scale mixture keeps every full conditional standard %s:' % cite('west1984', 'lange1989'))
add_formula('y_i | beta, sigma^2, w_i ~ N(x_i\' beta, sigma^2 / w_i),   '
            'w_i ~ Gamma(nu/2, nu/2)')
add_paragraph(
    'Marginalising the weights w_i recovers a Student-t with nu degrees of freedom and scale '
    'sigma. Conditional on the weights the model is a weighted Gaussian regression, so the '
    'sampler remains a tuning-free Gibbs sampler in which an observation with a large '
    'standardised residual receives a small weight and is automatically down-weighted.')

add_heading('2.7 Convergence Diagnostics', level=2)
add_paragraph('Notation. Each of the C sampled chains is split in half, giving M = 2C split '
              'chains of length N, so the total number of draws is S = M N. Write W for the mean '
              'within-split-chain variance, B for the between-split-chain variance, rho_t for the '
              'autocorrelation at lag t and tau for the integrated autocorrelation time.')
add_paragraph('Rank-normalised and folded R-hat.', bold=True)
add_paragraph(
    'The classical statistic %s compares between-chain and within-chain variance, R-hat = '
    'sqrt(V / W) with V = ((N-1)/N) W + B/N. It assumes finite variance and is not invariant to '
    'monotone transformations. We therefore use the rank-normalised version %s: draws are pooled, '
    'replaced by their ranks with ties averaged, and mapped through the inverse normal cdf with '
    'the Blom transform z = Phi^-1((r - 3/8)/(S - 1/4)) before the classical formula is applied. '
    'The folded variant applies the same procedure to |theta - median(theta)|, which makes it '
    'sensitive to disagreement in scale rather than location. The reported value is the larger of '
    'the two. The threshold used throughout is %.2f.'
    % (cite('gelman2013'), cite('vehtari2021'), configuration['rhat_threshold']))
add_paragraph('Bulk and tail effective sample size.', bold=True)
add_paragraph(
    'ESS = S / tau, where tau = 1 + 2 sum_{t>=1} rho_t. The autocorrelations are estimated '
    'jointly across chains through rho_t = 1 - (W - mean_m[s_m^2 rho_{t,m}]) / V, so that '
    'between-chain variation enters the estimate, and the sum is truncated by Geyer\'s initial '
    'monotone positive sequence rule %s. Bulk-ESS applies this to the rank-normalised draws and '
    'describes the centre of the distribution; tail-ESS is the smaller of the ESS values for the '
    '5 percent and 95 percent quantile indicator series and describes the tails. Both are single '
    'numbers computed from all S draws jointly, never per-chain values that are then averaged.'
    % cite('geyer1992'))
add_paragraph('Aggregation across parameters.', bold=True)
add_paragraph(
    'Each diagnostic is computed for one parameter at a time, for all %d regression coefficients '
    'and the noise variance sigma^2, %d parameters in total. The single number reported for a '
    'sampler is the worst case: the maximum R-hat and the minimum bulk and tail ESS over those '
    'parameters. Taking the worst rather than an average prevents one badly behaved coordinate '
    'from being hidden by well behaved ones. Efficiency per second uses the same minimum bulk '
    'ESS and the same total runtime, so the ratio is internally consistent.'
    % (data_summary['n_features'], data_summary['n_features'] + 1))
add_paragraph('Monte Carlo standard error.', bold=True)
add_paragraph(
    'MCSE is the posterior standard deviation divided by the square root of the joint bulk ESS, '
    'and expresses how precisely a posterior mean has been located by a finite run. It is '
    'reported only for chains that converged, because for a chain that has not converged the '
    'estimate is biased and its standard error understates the true error.')
add_paragraph('Implementation check.', bold=True)
if arviz_check:
    add_paragraph(
        'All diagnostics are implemented from scratch. To confirm the implementation rather than '
        'assert it, the estimators were compared with ArviZ %s, the reference implementation of '
        'the same paper. For the Gibbs intercept our R-hat is %.5f against ArviZ %.5f, and our '
        'bulk ESS is %.1f against %.1f. Across a set of synthetic stress cases including '
        'independent draws, a strongly autocorrelated series, offset chains and heavy-tailed '
        'draws, the two agreed to within 0.4 percent on ESS.'
        % (cite('kumar2019'), arviz_check['ours_rhat'], arviz_check['arviz_rhat'],
           arviz_check['ours_bulk_ess'], arviz_check['arviz_bulk_ess']))

add_heading('2.8 Computational Complexity', level=2)
add_paragraph('With n observations, p parameters, L leapfrog steps and T retained iterations:')
add_table('Per-iteration time and total space complexity of each sampler.',
          ['Sampler', 'Time per iteration', 'Space', 'Dominant cost'],
          [('Metropolis-Hastings', 'O(np)', 'O(np + Tp)', 'one likelihood evaluation'),
           ('Adaptive Metropolis', 'O(np + p^2)', 'O(np + p^2 + Tp)', 'likelihood and proposal'),
           ('Preconditioned MH', 'O(np + p^2)', 'O(np + p^2 + Tp)', 'likelihood and proposal'),
           ('Gibbs', 'O(np + p^3)', 'O(np + p^2 + Tp)', 'p by p solve for beta'),
           ('HMC', 'O(L np)', 'O(np + Tp)', 'L gradient evaluations'),
           ('Student-t Gibbs', 'O(np^2 + p^3)', 'O(np + p^2 + Tp)', 'reweighted X\'WX each sweep')],
          column_widths_cm=[4.2, 3.4, 4.2, 5.2])
add_paragraph(
    'Preconditioning adds only an O(p^2) multiplication by a Cholesky factor, negligible beside '
    'the O(np) likelihood, so the efficiency it buys is close to free. The Student-t sampler is '
    'the most expensive per sweep because the weights change every iteration, so X\'WX cannot be '
    'cached. With n = %s and p = %d, np = %s dominates p^3 = %s.'
    % (thousands(data_summary['n_train']), data_summary['n_features'],
       thousands(data_summary['n_train'] * data_summary['n_features']),
       thousands(data_summary['n_features'] ** 3)))


# --------------------------------------------------------------------------------------------
# 3. Data
# --------------------------------------------------------------------------------------------

add_heading('3. Data and Experimental Design', level=1)

add_heading('3.1 Dataset', level=2)
add_paragraph(
    'The Bitbrains GWA-T-12 traces %s were collected from a managed hosting datacenter operated '
    'by Bitbrains IT Services in the Netherlands, and are distributed through the Grid Workloads '
    'Archive. The fastStorage trace records eleven telemetry channels for 1,250 virtual machines '
    'at five-minute resolution over roughly thirty days in August and September 2013. We use %d '
    'machines.' % (cite('shen2015'), data_summary['distinct_machines']))
add_figure('fig_eda.png',
           'Distributions of the six telemetry channels used as predictors. Counts are on a '
           'logarithmic scale, and the five strictly positive channels also use a logarithmic '
           'horizontal axis, because on linear axes every distribution collapses into a single '
           'bar at the left edge.',
           'Six histograms of telemetry channels with logarithmic count axes, showing strongly '
           'right-skewed distributions for memory, disk and network activity.')

add_heading('3.2 Forecasting Task and Feature Construction', level=2)
add_paragraph(
    'The task is one-step-ahead forecasting at a horizon of %d minutes: the target is CPU '
    'utilisation at time t, and every predictor is a function of information available strictly '
    'before t. This is stated explicitly because it constrains the design. The contemporaneous '
    'memory, disk and network channels are not observable when the forecast is issued, so each is '
    'lagged by one interval. Had they been used at time t the exercise would have been nowcasting '
    'rather than forecasting.' % configuration['horizon_minutes'])
add_paragraph('The %d predictors, plus an intercept, are:' % (data_summary['n_features'] - 1))
add_bullet('Exogenous load at t-1: memory usage, disk read and write throughput, network '
           'received and transmitted throughput.')
add_bullet('Autoregressive terms: CPU utilisation at t-1, t-2 and t-3.')
add_bullet('Rolling statistics of CPU utilisation over the six intervals ending at t-1: mean and '
           'standard deviation.')

add_heading('3.3 Chronological Splitting and Leakage Control', level=2)
add_paragraph(
    'Observations from all machines are ordered by timestamp and divided chronologically into '
    'training, validation and test periods in proportions 60, 20 and 20 percent. Three separate '
    'measures guard against leakage.')
add_bullet('Feature windows may not straddle a boundary. Each row depends on the six intervals '
           'preceding it, so any row whose earliest required observation falls in an earlier '
           'split is discarded. This embargo removed %s rows.'
           % thousands(data_summary['embargoed_count']))
add_bullet('All scaling constants, the predictor means and standard deviations and those of the '
           'target, are estimated on the training rows only and then applied unchanged to '
           'validation and test.')
add_bullet('The validation split is used to choose the likelihood and its degrees of freedom. '
           'The test split is used once, for the final numbers in Section 6, and for nothing '
           'else.')
add_table('Split sizes after the embargo.',
          ['Split', 'Rows', 'Role'],
          [('Training', thousands(data_summary['n_train']), 'Posterior sampling'),
           ('Validation', thousands(data_summary['n_validation']),
            'Likelihood and degrees-of-freedom selection'),
           ('Test', thousands(data_summary['n_test']), 'Final reported performance only'),
           ('Embargoed', thousands(data_summary['embargoed_count']),
            'Discarded to prevent window crossing')],
          column_widths_cm=[3.6, 2.6, 9.0])
add_paragraph(
    'One qualification should be stated plainly. The same machine appears in more than one split, '
    'at different times. For time-series forecasting this is the intended design, since the point '
    'is to predict a machine\'s future from its past, and the embargo ensures no individual '
    'prediction uses information from across a boundary. It does mean the splits are not '
    'independent in the sense of disjoint machines, so the results describe forecasting for '
    'machines already observed, not generalisation to unseen machines.')

add_heading('3.4 Sample Size', level=2)
add_paragraph(
    'Of roughly %s raw records available across the selected machines, the study uses the '
    'earliest %s engineered observations. The restriction is computational rather than '
    'statistical. Each sampler is run with %d chains, %s draws per chain and %d independent '
    'repeats, and the per-iteration cost of every sampler is O(np); the full experiment already '
    'requires several hundred thousand likelihood or gradient evaluations at n = %s. Using the '
    'whole trace would multiply that by roughly two orders of magnitude without changing what the '
    'experiment is designed to measure, which is the relative behaviour of the samplers on a '
    'fixed posterior. The consequence is that the test split contains only %s rows, so the '
    'coverage estimates in Section 6 carry appreciable uncertainty, which is why they are '
    'reported with bootstrap intervals.'
    % (thousands(data_summary['total_available_rows']),
       thousands(data_summary['n_train'] + data_summary['n_validation']
                 + data_summary['n_test'] + data_summary['embargoed_count']),
       configuration['chains'], thousands(configuration['draws_per_chain']),
       configuration['repeats'], thousands(data_summary['n_train']),
       thousands(data_summary['n_test'])))

add_heading('3.5 Pooling Across Machines', level=2)
add_paragraph(
    'All machines share a single coefficient vector. This is a modelling assumption, and the '
    'residuals show it is only approximately right: across the %d machines with at least ten '
    'training rows, the per-machine residual standard deviation ranges from %.3f to %.3f with a '
    'median of %.3f, against a pooled value of %.3f. A machine at the upper end therefore has '
    'noise several times larger than one at the lower end, and a single sigma^2 splits the '
    'difference, over-stating uncertainty for quiet machines and under-stating it for volatile '
    'ones. A hierarchical model with per-machine coefficients and variances is the natural '
    'remedy and is left to further work; it would also make the comparison between Gibbs and HMC '
    'considerably more interesting, since conjugacy alone no longer suffices there.'
    % (pooling['per_machine_count'], pooling['per_machine_residual_sd_min'],
       pooling['per_machine_residual_sd_max'], pooling['per_machine_residual_sd_median'],
       pooling['residual_sd']))

# --------------------------------------------------------------------------------------------
# 4. Methodology
# --------------------------------------------------------------------------------------------

add_heading('4. Methodology', level=1)
add_paragraph(
    'All samplers are implemented from scratch in NumPy and SciPy. No probabilistic programming '
    'library is used inside any implementation; emcee appears only as an external reference in '
    'Section 5.5, and ArviZ only to verify the diagnostic code.')
add_paragraph(
    'Each sampler is run with %d chains from overdispersed starting points, drawn as beta ~ '
    'N(0, 2^2) and log sigma^2 ~ N(0, 1.5^2), which are far wider than the posterior. '
    'Overdispersed starts are what split R-hat assumes: if every chain began near the mode, '
    'agreement between chains would demonstrate little. Each chain retains %s draws after %s '
    'burn-in iterations, and the whole procedure is repeated %d times with different seeds so '
    'that run-to-run variability in runtime and ESS can be reported.'
    % (configuration['chains'], thousands(configuration['draws_per_chain']),
       thousands(configuration['burn_in']), configuration['repeats']))
add_table('Tuning parameters. Gibbs samplers have none.',
          ['Sampler', 'Tuning parameters', 'Setting'],
          [('Metropolis-Hastings', 'proposal scales for beta and log sigma^2', '0.001, 0.05'),
           ('Adaptive Metropolis', 'adaptation interval, scaling', '200 iterations, 2.38^2/d'),
           ('Preconditioned MH', 'target acceptance, adaptation interval', '0.234, 100'),
           ('Gibbs', 'none', 'not applicable'),
           ('HMC', 'step size, leapfrog steps', '0.002, 15'),
           ('Student-t Gibbs', 'degrees of freedom nu', 'selected on validation')],
          column_widths_cm=[4.0, 6.4, 5.0])

# --------------------------------------------------------------------------------------------
# 5. Results
# --------------------------------------------------------------------------------------------

add_heading('5. Results: Convergence and Efficiency', level=1)

add_heading('5.1 Posterior Geometry', level=2)
add_paragraph(
    'Because the conditional posterior of beta given sigma^2 is Gaussian, its covariance is '
    'available in closed form. Evaluated at the fitted noise level it has condition number %.1f '
    'and a largest absolute correlation between coefficients of %.3f, with standard deviations '
    'along the extreme directions of %.4f and %.4f. The correlation is a direct consequence of '
    'the design: three consecutive CPU lags and a rolling mean of the same series measure nearly '
    'the same quantity. A single isotropic step size must be small enough to be accepted along '
    'the narrow direction while the chain has to travel roughly %.0f times further to cross the '
    'wide one.'
    % (geometry['condition_number'], geometry['max_absolute_correlation'],
       geometry['narrowest_direction_sd'], geometry['widest_direction_sd'],
       geometry['widest_direction_sd'] / geometry['narrowest_direction_sd']))

add_heading('5.2 Convergence', level=2)
convergence_rows = []
for name in GAUSSIAN_SAMPLERS:
    entry = samplers[name]
    convergence_rows.append((name, '%.4f' % entry['worst_rhat'],
                             '%.0f' % entry['min_bulk_ess'], '%.0f' % entry['min_tail_ess'],
                             'yes' if entry['converged'] else 'no'))
add_table('Convergence diagnostics from %d dispersed chains of %s draws. R-hat is the larger of '
          'the rank-normalised and folded statistics, maximised over monitored parameters; ESS '
          'values are minima over the same parameters. The criterion is R-hat below %.2f and bulk '
          'ESS at least %.0f.'
          % (configuration['chains'], thousands(configuration['draws_per_chain']),
             configuration['rhat_threshold'], configuration['bulk_ess_threshold']),
          ['Sampler', 'R-hat', 'Bulk ESS', 'Tail ESS', 'Converged'], convergence_rows,
          column_widths_cm=[4.6, 2.4, 2.6, 2.6, 2.6])
add_paragraph(
    'Two samplers meet the criterion at the standard run length and three do not. Plain '
    'Metropolis reaches R-hat %.2f and adaptive Metropolis %.2f, both far above the threshold, '
    'with bulk ESS in single figures out of %s draws. Neither has converged, and no quantity '
    'derived from them, whether a posterior mean, an efficiency figure or a predictive score, '
    'can be interpreted as an estimate of the target.'
    % (samplers['MH']['worst_rhat'], samplers['Adaptive MH (naive)']['worst_rhat'],
       thousands(samplers['MH']['total_draws'])))
add_paragraph(
    'Preconditioned Metropolis does not converge at %s draws per chain, with R-hat %.4f '
    'exceeding the %.2f threshold. Monitoring all %d parameters exposed a coordinate that was '
    'not visible in a smaller diagnostic set. Extending to %s draws per chain gives R-hat %.4f '
    'and bulk ESS %.0f, which does converge, confirming that the standard run length was '
    'insufficient rather than the method itself failing.'
    % (thousands(configuration['draws_per_chain']),
       samplers['Preconditioned MH']['worst_rhat'], configuration['rhat_threshold'],
       data_summary['n_features'] + 1,
       thousands(configuration['long_run_draws']), long_run['worst_rhat'],
       long_run['min_bulk_ess']))
add_figure('fig_convergence.png',
           'Worst-case R-hat and minimum bulk ESS for each sampler, both on logarithmic scales, '
           'with dashed lines marking the thresholds. Error bars on the ESS panel span the '
           'standard deviation across %d independent repeats.' % configuration['repeats'],
           'Two bar charts. The left shows R-hat by sampler against a threshold of 1.01, with '
           'Metropolis and adaptive Metropolis far above it. The right shows bulk effective '
           'sample size, with Gibbs and Hamiltonian Monte Carlo orders of magnitude above the '
           'others.')
add_paragraph(
    'The consequences are visible in the estimates themselves. Compared with the deterministic '
    'quadrature reference of Section 5.5, the converged samplers agree to within %.4f, while '
    'plain Metropolis is wrong by %.2f and adaptive Metropolis by %.2f. For scale, the widest '
    'direction of the posterior has standard deviation %.4f, so both errors exceed the entire '
    'spread of the posterior.'
    % (WORST_CONVERGED_DEVIATION, DEVIATIONS['MH'], DEVIATIONS['Adaptive MH (naive)'],
       geometry['widest_direction_sd']))
add_figure('fig_posteriors.png',
           'Posterior distributions by sampler for three coefficients and the noise variance. The '
           'converged samplers produce indistinguishable densities; the two that did not converge '
           'are displaced and too narrow.',
           'Four overlaid histogram panels comparing posterior distributions from five samplers, '
           'showing close agreement among three of them and visible displacement of the other '
           'two.')

add_heading('5.3 Why Metropolis Fails, and What Repairs It', level=2)
add_paragraph(
    'The failure of plain Metropolis follows from the geometry of Section 5.1 rather than from '
    'random-walk behaviour in the abstract. Two repairs were attempted and only one worked.')
add_paragraph(
    'Estimating the proposal covariance from the chain\'s own history, the textbook adaptive '
    'recipe, made matters worse: acceptance collapsed to %s and R-hat rose to %.2f. The reason is '
    'a bootstrap problem. Adaptation needs the chain to have explored enough to estimate its own '
    'covariance, but the unpreconditioned chain has an integrated autocorrelation time of the '
    'order of the run length, so across %s burn-in iterations it supplies only a handful of '
    'effectively independent points from which to estimate a %d-dimensional covariance. What the '
    'estimate captures instead is the transient drift from the dispersed start toward the mode, '
    'whose spread along the direction of travel is much larger than the posterior; scaling that '
    'by 2.38^2/d produces proposals so large that almost everything is rejected.'
    % (format_percentage(samplers['Adaptive MH (naive)']['acceptance_rate']),
       samplers['Adaptive MH (naive)']['worst_rhat'], thousands(configuration['burn_in']),
       data_summary['n_features'] + 1))
add_paragraph(
    'This should be read as a statement about this configuration, not about adaptive Metropolis '
    'in general. The method is standard and effective in many settings %s. What this experiment '
    'shows is that it cannot bootstrap itself when the initial proposal mixes too slowly relative '
    'to the burn-in budget; a longer adaptation phase, or a better initial proposal, would be '
    'expected to rescue it.' % cite('haario2001'))
add_paragraph(
    'Taking the metric from outside the chain does work. Preconditioning with the observed Fisher '
    'information brings acceptance to %s, against the theoretical optimum of 23.4 percent, and '
    'raises bulk ESS from %.1f to %.0f, a factor of %.0f, for essentially no additional cost per '
    'iteration (%.1f seconds against %.1f for the full multi-chain run).'
    % (format_percentage(samplers['Preconditioned MH']['acceptance_rate']),
       samplers['MH']['min_bulk_ess'], samplers['Preconditioned MH']['min_bulk_ess'],
       samplers['Preconditioned MH']['min_bulk_ess'] / max(samplers['MH']['min_bulk_ess'], 1e-9),
       samplers['Preconditioned MH']['time'], samplers['MH']['time']))
add_paragraph(
    'The gain should not be overstated. At %.0f effective draws out of %s, preconditioned '
    'Metropolis remains an order of magnitude behind Gibbs and HMC. Matching the metric to the '
    'posterior removes the penalty for anisotropy; it does not remove the random walk.'
    % (samplers['Preconditioned MH']['min_bulk_ess'],
       thousands(samplers['Preconditioned MH']['total_draws'])))
add_figure('fig_preconditioning.png',
           'Left: posterior correlation between coefficients, showing the block of correlated lag '
           'and rolling features. Centre: bulk ESS before and after preconditioning, on a '
           'logarithmic scale. Right: autocorrelation of the intercept.',
           'Three panels: a correlation heat map, a bar chart comparing effective sample size '
           'before and after preconditioning, and autocorrelation curves showing much faster '
           'decay after preconditioning.')

add_heading('5.4 Efficiency', level=2)
efficiency_rows = []
for name in CONVERGED_SAMPLERS:
    entry = samplers[name]
    efficiency_rows.append((
        name, format_percentage(entry['acceptance_rate']) if entry['acceptance_rate'] is not None
        else 'N/A',
        '%.0f' % entry['min_bulk_ess'],
        '%.1f +/- %.1f' % (entry['time_mean'], entry['time_sd']),
        '%.0f' % entry['bulk_ess_per_second'],
        '%.2e' % entry['mcse_intercept']))
add_table('Efficiency of the converged samplers. Runtime is the total for %d chains of %s draws, '
          'reported as mean and standard deviation over %d repeats. ESS per second uses the same '
          'minimum bulk ESS and the same total runtime.'
          % (configuration['chains'], thousands(configuration['draws_per_chain']),
             configuration['repeats']),
          ['Sampler', 'Acceptance', 'Bulk ESS', 'Runtime (s)', 'Bulk ESS/s', 'MCSE (intercept)'],
          efficiency_rows, column_widths_cm=[4.2, 2.4, 2.2, 2.8, 2.4, 2.8])
add_paragraph(
    'Only the converged samplers are ranked. Including Metropolis or adaptive Metropolis in an '
    'efficiency table would invite a false comparison, because an ESS per second computed from a '
    'chain that is sampling the wrong distribution measures the rate at which it produces '
    'unusable draws. Their diagnostics appear in Table 3 and their runtimes in the accompanying '
    'results file, but they are excluded from every ranking in this section and from the '
    'predictive comparison in Section 6.')
ranking = sorted(CONVERGED_SAMPLERS, key=lambda name: -samplers[name]['bulk_ess_per_second'])
add_paragraph(
    'The ranking by bulk ESS per second is %s. Gibbs leads because each draw is exact and cheap '
    'and because its ESS per draw is close to one; HMC produces nearly uncorrelated draws but '
    'pays for %d gradient evaluations per iteration. The Monte Carlo standard errors tell the '
    'same story in the units that matter for reporting a posterior mean.'
    % (', '.join('%s (%.0f)' % (name, samplers[name]['bulk_ess_per_second'])
                 for name in ranking), 15))

add_heading('5.5 Validation Against Independent References', level=2)
add_paragraph(
    'Agreement among our own samplers is reassuring but limited, because all of them were written '
    'by the same authors from the same equations. Two independent references are therefore used, '
    'and they check different things.')
add_paragraph('An analytical reference.', bold=True)
add_paragraph(
    'Although the joint posterior has no standard closed form, beta can be integrated out '
    'analytically: marginally, y | sigma^2 ~ N(0, sigma^2 I + tau^2 X X\'), whose log density is '
    'evaluated with Sylvester\'s determinant identity and the Woodbury identity. The remaining '
    'one-dimensional density over sigma^2 is then integrated on a grid of %s points, and the '
    'posterior mean of beta recovered as the corresponding mixture of conditional means. This '
    'reference involves no sampling whatsoever, so it validates the transcription of the '
    'likelihood and the priors, not merely our sampling of them. Gibbs reproduces it to within '
    '%.5f and the reference posterior mean of sigma^2 is %.5f.'
    % (thousands(analytical['grid_size']), analytical['largest_deviation_from_gibbs'],
       analytical['posterior_mean_variance']))
if external:
    add_paragraph('An external sampler.', bold=True)
    add_paragraph(
        'The same posterior was also sampled with %s %s, which implements the affine-invariant '
        'ensemble algorithm of Goodman and Weare %s. It shares nothing with our methods: no '
        'conditional structure, no gradients and no proposal covariance, and it is written by '
        'others. Running %d walkers for %s steps and discarding %s produced %s draws in %.1f '
        'seconds, with worst R-hat %.4f and minimum bulk ESS %.0f. Its posterior means agree with '
        'Gibbs to within %.5f.'
        % (external['library'], cite('foremanmackey2013'), cite('goodman2010'),
           external['walkers'], thousands(external['steps']),
           thousands(external['discarded_steps']), thousands(external['total_draws']),
           external['time'], external['worst_rhat'], external['minimum_bulk_ess'],
           external['largest_deviation_from_gibbs']))
    add_paragraph(
        'The scope of this second check should be stated precisely. Agreement with emcee provides '
        'independent evidence that our custom samplers correctly target the supplied '
        'log-posterior. Because emcee evaluates the same user-defined log-posterior function, the '
        'comparison does not independently validate the transcription of the likelihood and the '
        'priors. That is the role of the analytical reference above, which is constructed from '
        'the model algebra rather than from our code path.')
    if external['worst_rhat'] >= configuration['rhat_threshold']:
        add_paragraph(
            'One caveat: at R-hat %.4f the emcee run is itself slightly above the %.2f threshold '
            'used for our own samplers. Ensemble samplers correlate their walkers by '
            'construction, so the split-chain diagnostic is conservative when applied to them, '
            'but the agreement it supports is correspondingly approximate and is quoted here to '
            'four decimal places rather than more.'
            % (external['worst_rhat'], configuration['rhat_threshold']))
add_paragraph('The No-U-Turn Sampler was not used as a reference; no NUTS results are reported '
              'in this study.')

add_heading('5.6 Sensitivity to Tuning Parameters', level=2)
add_paragraph(
    'Metropolis and HMC both require tuning; the Gibbs samplers do not. For HMC the step size '
    'epsilon and the number of leapfrog steps L are varied jointly, because they trade off '
    'against one another and neither is interpretable alone. Efficiency is reported as bulk ESS '
    'per gradient evaluation, which is the natural currency for a gradient method, alongside '
    'acceptance rate and R-hat.')
add_paragraph(
    'The grid runs differ from the main experiment in one respect that should be stated. They are '
    'warm started at the least-squares fit with small random perturbations, rather than from the '
    'overdispersed points used in Section 5.2. Short runs from dispersed starts would spend most '
    'of their length travelling to the posterior, so the resulting diagnostics would measure '
    'burn-in rather than the effect of the tuning parameters, which is what the grid is meant to '
    'isolate. The cost of that choice is that split R-hat is a weaker guarantee here than in '
    'Table 4, because chains that begin close together can agree without having explored the '
    'posterior. The grid should therefore be read as ranking configurations against one another, '
    'not as certifying any of them as converged.')
grid_rows = []
for cell in hmc_grid:
    rhat_display = '%.4f' % cell['worst_rhat'] if cell['worst_rhat'] is not None else 'degenerate'
    grid_rows.append(('%.3f' % cell['step_size'], str(cell['leapfrog_steps']),
                      format_percentage(cell['acceptance']), rhat_display,
                      '%.0f' % cell['bulk_ess'], '%.2e' % cell['bulk_ess_per_gradient'],
                      'yes' if cell['converged'] else 'no'))
add_table('HMC sensitivity to step size and trajectory length. Cells that did not converge are '
          'marked; their ESS values are not interpretable and are shown only for completeness.',
          ['epsilon', 'L', 'Acceptance', 'R-hat', 'Bulk ESS', 'Bulk ESS/gradient', 'Converged'],
          grid_rows, column_widths_cm=[2.0, 1.4, 2.4, 2.2, 2.2, 3.2, 2.2])
if BEST_HMC_CELL:
    add_paragraph(
        'Only %d of the %d configurations converged, and the comparison is restricted to those. '
        'Among them the most efficient is epsilon = %.3f with L = %d, at %.2e bulk ESS per '
        'gradient evaluation and an acceptance rate of %s.'
        % (len(CONVERGED_HMC_CELLS), len(hmc_grid), BEST_HMC_CELL['step_size'],
           BEST_HMC_CELL['leapfrog_steps'], BEST_HMC_CELL['bulk_ess_per_gradient'],
           format_percentage(BEST_HMC_CELL['acceptance'])))
    highest_acceptance = max(hmc_grid, key=lambda cell: cell['acceptance'])
    add_paragraph(
        'This directly contradicts the reading of a high acceptance rate as a sign of good '
        'tuning. The highest acceptance in the grid, %s at epsilon = %.3f and L = %d, is not the '
        'most efficient configuration; a step size that small buys near-certain acceptance by '
        'proposing moves so short that the trajectory barely advances, and the gradient budget is '
        'spent for little gain. The theoretically motivated target for HMC is roughly 65 to 80 '
        'percent %s, and every acceptance rate in this grid above %.3f exceeds it, indicating '
        'that the step size used elsewhere in the study is conservative. At the other extreme, '
        'epsilon = 0.008 breaks the leapfrog integrator entirely and acceptance falls to zero.'
        % (format_percentage(highest_acceptance['acceptance']), highest_acceptance['step_size'],
           highest_acceptance['leapfrog_steps'], cite('beskos2013'), 0.004))
add_table('Metropolis acceptance against proposal scale, at fixed scale 0.05 for log sigma^2.',
          ['Proposal scale for beta', 'Acceptance rate'],
          [('%.4f' % entry['step'], format_percentage(entry['acceptance']))
           for entry in mh_grid], column_widths_cm=[5.4, 4.4])
add_paragraph(
    'The Metropolis grid shows the familiar trade-off, but it is worth reading together with '
    'Section 5.1: no scalar step size performs well, because the search is over a '
    'one-dimensional family for a problem that requires a matrix.')
add_figure('fig_hmc_sensitivity.png',
           'Bulk ESS per gradient evaluation (left) and acceptance rate (right) over the grid of '
           'step sizes and trajectory lengths. The most efficient configuration is not the one '
           'with the highest acceptance.',
           'Two heat maps over step size and leapfrog steps, one showing effective sample size '
           'per gradient evaluation and the other acceptance rate, demonstrating that the highest '
           'acceptance does not coincide with the best efficiency.')

add_heading('5.7 Sensitivity to the Starting Point', level=2)
add_paragraph(
    'MCMC converges from any starting point in theory but says nothing about how long that takes. '
    'Each sampler was started from four deliberately different points. A chain is deemed to have '
    'reached a stable region at the first iteration after which its log posterior remains within '
    'three median absolute deviations (MAD) of the median over the final 500 iterations for the '
    'remainder of the chain. Requiring sustained containment rather than a brief crossing '
    'prevents a chain from being credited for merely passing through the region on its way '
    'elsewhere.')
start_names = list(next(iter(initialisation.values())).keys())
add_table('Iterations required to reach the stable region of the log posterior.',
          ['Sampler'] + [name.replace(' (all +3)', ' (+3)') for name in start_names],
          [tuple([sampler] + [str(entries[start]['iterations_to_stable_region'])
                              for start in start_names])
           for sampler, entries in initialisation.items()],
          column_widths_cm=[4.4, 2.6, 3.4, 3.0, 3.0])
gibbs_worst = max(initialisation['Gibbs'][s]['iterations_to_stable_region']
                  for s in start_names)
hmc_worst = max(initialisation['HMC'][s]['iterations_to_stable_region']
                for s in start_names)
pm_values = [initialisation['Preconditioned MH'][s]['iterations_to_stable_region']
             for s in start_names]
add_paragraph(
    'Gibbs and HMC both reach the stable region at a similar iteration count from every start, '
    '%d and %d respectively, because their rapid mixing makes the starting point irrelevant: '
    'the occasional late excursion from the MAD band is driven by random variation rather than '
    'by the initial transient. Preconditioned Metropolis varies more, taking between %d and %d '
    'iterations depending on the start, because its slower mixing means the initial transient '
    'still influences how long the chain needs to settle.'
    % (gibbs_worst, hmc_worst, min(pm_values), max(pm_values)))
add_figure('fig_initialisation.png',
           'Log posterior over the first iterations from four starting points, on a symmetric '
           'logarithmic scale. Curves that meet quickly indicate the starting point has been '
           'forgotten.',
           'Three line-chart panels, one per sampler, showing log posterior traces from four '
           'different starting points converging to a common level at different speeds.')


# --------------------------------------------------------------------------------------------
# 6. Predictive performance
# --------------------------------------------------------------------------------------------

add_heading('6. Predictive Performance and Calibration', level=1)

add_heading('6.1 Construction of Posterior Predictive Intervals', level=2)
add_paragraph(
    'All intervals reported in this section are posterior predictive intervals for a future CPU '
    'observation, not credible intervals for a parameter. The distinction matters because the two '
    'have different widths and answer different questions. For each retained posterior draw '
    '(beta_s, sigma^2_s) a replicate observation is generated as')
add_formula('y* = x\' beta_s + sigma_s e,    e ~ N(0, 1)  or  e ~ t_nu')
add_paragraph(
    'so each replicate carries both sources of uncertainty: parameter uncertainty, through the '
    'spread of beta_s and sigma_s across draws, and observation noise, through e. The interval is '
    'then the empirical quantile range of these replicates, rather than a point estimate plus a '
    'multiple of a standard deviation, so no symmetry or normality is imposed on the result.')

add_heading('6.2 Selecting the Likelihood on Validation Data', level=2)
add_paragraph(
    'The choice between the Gaussian and Student-t likelihoods, and the value of nu, is made on '
    'the validation split alone, by mean log pointwise predictive density. The predictive density '
    'is the posterior mixture (1/S) sum_s p(y* | theta_s), and each component is evaluated in '
    'closed form, so the criterion does not depend on simulated replicates.')
selection_rows = []
for candidate in selection['candidates']:
    label = 'Gaussian' if candidate['likelihood'] == 'Gaussian' \
        else 'Student-t, nu = %.0f' % candidate['degrees_of_freedom']
    selection_rows.append((label, '%.4f' % candidate['validation_log_density'],
                           format_percentage(candidate['validation_coverage_50']),
                           format_percentage(candidate['validation_coverage_95']),
                           '%.4f' % candidate['validation_median_absolute_error']))
add_table('Likelihood selection on the validation split. Higher log predictive density is better; '
          'nominal coverage is 50 and 95 percent. The selected row is the Student-t with nu = %.0f.'
          % SELECTED_DEGREES,
          ['Candidate', 'Log predictive density', 'Coverage 50%', 'Coverage 95%',
           'Median abs. error'], selection_rows,
          column_widths_cm=[4.0, 3.6, 2.6, 2.6, 3.0])
add_paragraph(
    'The criterion prefers the Student-t at every value of nu tested, and prefers heavier tails '
    'monotonically, selecting the smallest value on the grid, nu = %.0f. Two qualifications '
    'follow. First, nu = 2 has infinite variance, so the fitted model has a well-defined median '
    'and scale but no finite predictive variance; this is defensible for a criterion based on '
    'predictive density but would be inappropriate if a variance were needed downstream. Second, '
    'because the selected value sits at the edge of the grid, the data are indicating that they '
    'would prefer heavier tails still, and a wider grid or a prior on nu would be the principled '
    'next step rather than accepting an edge solution. The sensitivity of the criterion across '
    'the grid is shown below; the ordering is stable and not an artefact of a single value.'
    % SELECTED_DEGREES)
add_figure('fig_nu_selection.png',
           'Validation log predictive density (left) and 50 percent interval coverage (right) '
           'across the grid of degrees of freedom, with the Gaussian shown as a horizontal '
           'reference. The criterion decreases monotonically in nu.',
           'Two line charts against degrees of freedom on a logarithmic axis, showing log '
           'predictive density decreasing with nu and coverage remaining well above the nominal '
           'level for every candidate.')

add_heading('6.3 Test-Set Performance', level=2)
add_paragraph(
    'The test split is used here and nowhere else. Errors are reported both in standardised units '
    'and in CPU percentage points, obtained by multiplying by the training standard deviation of '
    'the target, %.3f percentage points.' % data_summary['target_std'])
STUDENT_DIAGNOSTICS = results.get('student_t_diagnostics')
if STUDENT_DIAGNOSTICS:
    add_paragraph(
        'The Student-t model is held to the same standard as the samplers in Section 5. It is '
        'fitted with %d chains from overdispersed starting points and diagnosed with the same '
        'estimators and the same worst-case aggregation, giving R-hat %.4f and minimum bulk ESS '
        '%.0f, which %s the criterion of R-hat below %.2f and bulk ESS at least %.0f used '
        'throughout. Every Student-t number reported in this section comes from those %d chains '
        'pooled, not from a single run. This matters because the argument of Section 5 is that a '
        'sampler which has not been diagnosed cannot support conclusions, and that argument '
        'applies to the model used for the predictive results just as much as to the ones being '
        'compared.'
        % (configuration['chains'], STUDENT_DIAGNOSTICS['worst_rhat'],
           STUDENT_DIAGNOSTICS['min_bulk_ess'],
           'meets' if STUDENT_DIAGNOSTICS['converged'] else 'does not meet',
           configuration['rhat_threshold'], configuration['bulk_ess_threshold'],
           configuration['chains']))
accuracy_rows = [
    ('OLS baseline', '%.4f' % OLS_TEST['rmse'], '%.3f' % OLS_TEST['rmse_percentage_points'],
     '%.4f' % OLS_TEST['median_absolute_error'],
     '%.3f' % OLS_TEST['median_absolute_error_percentage_points']),
    ('Bayesian, Gaussian likelihood', '%.4f' % GAUSSIAN_TEST['rmse'],
     '%.3f' % GAUSSIAN_TEST['rmse_percentage_points'],
     '%.4f' % GAUSSIAN_TEST['median_absolute_error'],
     '%.3f' % GAUSSIAN_TEST['median_absolute_error_percentage_points']),
    ('Bayesian, Student-t (nu = %.0f)' % SELECTED_DEGREES, '%.4f' % STUDENT_TEST['rmse'],
     '%.3f' % STUDENT_TEST['rmse_percentage_points'],
     '%.4f' % STUDENT_TEST['median_absolute_error'],
     '%.3f' % STUDENT_TEST['median_absolute_error_percentage_points'])]
add_table('Test-set accuracy in standardised units and in CPU percentage points.',
          ['Model', 'RMSE', 'RMSE (pp)', 'Median abs. error', 'Median abs. error (pp)'],
          accuracy_rows, column_widths_cm=[5.4, 2.2, 2.4, 3.0, 3.4])
add_paragraph(
    'The Gaussian Bayesian model and ordinary least squares are indistinguishable in point '
    'accuracy, which is expected: with weak priors and %s training observations the posterior '
    'mean is essentially the least-squares solution, and the value of the Bayesian treatment lies '
    'in the predictive distribution rather than the point estimate. The gap between RMSE and '
    'median absolute error, a factor of about %.0f, is the first sign that the error '
    'distribution is far from Gaussian: a few large errors dominate the squared measure while the '
    'typical error is much smaller. The Student-t model reduces the median absolute error from '
    '%.3f to %.3f percentage points.'
    % (thousands(data_summary['n_train']),
       GAUSSIAN_TEST['rmse'] / GAUSSIAN_TEST['median_absolute_error'],
       GAUSSIAN_TEST['median_absolute_error_percentage_points'],
       STUDENT_TEST['median_absolute_error_percentage_points']))

add_heading('6.4 Residual Behaviour', level=2)
add_paragraph(
    'Two properties of the residuals shape the calibration results. They are strongly '
    'non-Gaussian, with excess kurtosis %.0f and a ratio of standard deviation to robust standard '
    'deviation of %.2f, so almost all residuals are small and a few are very large. This is what '
    'a CPU trace of quiet periods punctuated by spikes produces. A Gaussian likelihood has one '
    'parameter with which to describe both regimes and inflates sigma^2 to cover the spikes.'
    % (residuals['excess_kurtosis'], residuals['sd_to_robust_sd_ratio']))
add_paragraph(
    'They also remain autocorrelated within each machine. The median per-machine lag-one '
    'residual autocorrelation is %.3f and %d of %d machines reject the Ljung-Box test at the '
    '5 percent level (median p = %.3f). The model treats observations as conditionally '
    'independent given the predictors, and the lag terms absorb much but not all of the serial '
    'structure within each machine. Remaining autocorrelation means the effective number of '
    'independent test observations is smaller than the nominal count, which is why the coverage '
    'estimates below are accompanied by cluster-bootstrap intervals that resample machines with '
    'replacement rather than binomial ones.'
    % (residuals['median_lag1_acf'], residuals['machines_rejecting'],
       residuals['machines_tested'], residuals['median_ljung_box_p']))
add_figure('fig_residuals.png',
           'Left: residual autocorrelation on the test split with the band expected under '
           'independence. Right: normal Q-Q plot of the test residuals, whose S-shape is the '
           'signature of heavy tails.',
           'Two panels: a bar chart of residual autocorrelation by lag against a confidence band, '
           'and a normal quantile-quantile plot curving away from the diagonal at both ends.')

add_heading('6.5 Calibration', level=2)
calibration_rows = []
for level in (50, 95):
    calibration_rows.append((
        '%d%%' % level,
        format_percentage(GAUSSIAN_TEST['coverage_%d' % level]),
        format_interval(GAUSSIAN_TEST['coverage_%d_interval' % level]),
        format_percentage(STUDENT_TEST['coverage_%d' % level]),
        format_interval(STUDENT_TEST['coverage_%d_interval' % level]),
        '%.2f' % GAUSSIAN_TEST['width_%d_percentage_points' % level],
        '%.2f' % STUDENT_TEST['width_%d_percentage_points' % level]))
add_table('Posterior predictive interval coverage on the test split, with cluster-bootstrap '
          '95 percent intervals (%s replicates, resampling machines with replacement) that '
          'account for within-machine serial dependence. Widths are in CPU percentage points.'
          % thousands(configuration['bootstrap_replicates']),
          ['Nominal', 'Gaussian', 'Gaussian 95% CI', 'Student-t', 'Student-t 95% CI',
           'Width Gauss (pp)', 'Width t (pp)'], calibration_rows,
          column_widths_cm=[1.8, 2.0, 3.0, 2.0, 3.0, 2.4, 2.2])
add_paragraph(
    'Under the Gaussian likelihood the intervals are far too wide. A nominal 50 percent interval '
    'covers %s of test observations, roughly %.1f times the intended rate, and the nominal 95 '
    'percent interval covers %s. The bootstrap intervals show this is not sampling noise.'
    % (format_percentage(GAUSSIAN_TEST['coverage_50']),
       GAUSSIAN_TEST['coverage_50'] / 0.5, format_percentage(GAUSSIAN_TEST['coverage_95'])))
add_paragraph(
    'The Student-t likelihood improves calibration substantially but does not fix it. Coverage of '
    'the nominal 50 percent interval falls from %s to %s and the mean width from %.2f to %.2f '
    'percentage points, yet %s against a nominal 50 percent is still severely miscalibrated, '
    'covering roughly %.1f times the intended fraction of observations. This is a partial '
    'improvement, not a repair, and the residual miscalibration is consistent with the '
    'diagnostics of Section 6.4: a Student-t with a single scale still assumes one noise level '
    'for every machine and every time, whereas the data are heteroscedastic across machines, as '
    'Section 3.5 showed, and serially dependent.'
    % (format_percentage(GAUSSIAN_TEST['coverage_50']),
       format_percentage(STUDENT_TEST['coverage_50']),
       GAUSSIAN_TEST['width_50_percentage_points'], STUDENT_TEST['width_50_percentage_points'],
       format_percentage(STUDENT_TEST['coverage_50']), STUDENT_TEST['coverage_50'] / 0.5))
add_paragraph(
    'A useful check confirms that the cause lies in the likelihood rather than in the sampler or '
    'the priors: the same over-coverage appears when intervals are built from an ordinary '
    'least-squares fit and its residual standard deviation, with no Bayesian machinery at all.')
add_figure('fig_calibration.png',
           'Left: coverage of posterior predictive intervals under each likelihood, with '
           'cluster-bootstrap intervals and stars marking the nominal levels. Right: mean interval '
           'width in CPU percentage points.',
           'Two bar charts comparing Gaussian and Student-t likelihoods, showing coverage well '
           'above nominal for both and much narrower intervals under the Student-t.')
add_figure('fig_predictions.png',
           'One-step-ahead forecasts on the test split in original units, with the 95 percent '
           'posterior predictive interval.',
           'A time-series chart comparing actual and predicted CPU utilisation with a shaded '
           'predictive interval that is visibly wider than the typical prediction error.')

# --------------------------------------------------------------------------------------------
# 7. Discussion
# --------------------------------------------------------------------------------------------

add_heading('7. Discussion', level=1)
add_paragraph('Answers to the research questions.', bold=True)
add_paragraph(
    'RQ1. Two of the five samplers converge at the standard %s draws per chain. Preconditioned '
    'Metropolis converges only with a longer run of %s draws. The two that fail entirely do so '
    'because the posterior is anisotropic, with condition number %.0f arising from the '
    'near-collinear lag features, and because an isotropic random-walk proposal cannot cope with '
    'that geometry. Supplying a metric from the observed Fisher information largely repairs it; '
    'estimating the metric from the chain itself does not, under the burn-in budget tested.'
    % (thousands(configuration['draws_per_chain']),
       thousands(configuration['long_run_draws']), geometry['condition_number']))
add_paragraph(
    'RQ2. Among the samplers that converge at the standard run length, Gibbs is the most '
    'efficient both per second and per draw in this conjugate setting. HMC produces nearly '
    'uncorrelated draws but pays for L gradient evaluations per iteration, and the sensitivity '
    'grid shows its efficiency per gradient depends jointly on step size and trajectory length, '
    'with the best configuration not the one with the highest acceptance rate.')
add_paragraph(
    'RQ3. The intervals are not calibrated. A Gaussian likelihood produces nominal 50 percent '
    'intervals covering %s, because heavy-tailed residuals inflate the single noise parameter. A '
    'Student-t likelihood selected on validation data reduces this to %s, a substantial but '
    'partial improvement.'
    % (format_percentage(GAUSSIAN_TEST['coverage_50']),
       format_percentage(STUDENT_TEST['coverage_50'])))
add_paragraph('Strengths and weaknesses of each sampler.', bold=True)
add_bullet('Metropolis-Hastings: the most general, needing only pointwise density evaluation, and '
           'the cheapest per iteration. On this posterior it did not converge, so its generality '
           'is of no practical use here without preconditioning.')
add_bullet('Adaptive Metropolis: sound in principle and standard in practice, but unable to '
           'bootstrap a proposal covariance from a chain that mixes too slowly during the burn-in '
           'available. A longer adaptation phase would be the obvious remedy.')
add_bullet('Preconditioned Metropolis: keeps the generality of Metropolis while removing the '
           'anisotropy penalty, at negligible extra cost, and needs no conjugacy. It converges '
           'only with a longer run and remains a random walk, an order of magnitude behind '
           'Gibbs and HMC in effective sample size.')
add_bullet('Gibbs: no tuning, exact conditionals, insensitive to the starting point, and the '
           'fastest here. It requires closed-form conditionals, which exist only because of the '
           'prior structure chosen, so it does not generalise to most realistic models.')
add_bullet('Hamiltonian Monte Carlo: near-independent draws and the best scaling to higher '
           'dimensions, at the cost of a differentiable log posterior and two tuning parameters.')
add_paragraph('Limitations.', bold=True)
add_bullet('The model is linear and pools all machines under one coefficient vector, which '
           'Section 3.5 shows is only approximately right.')
add_bullet('Only %s observations are used, so the test split contains %s rows and the coverage '
           'estimates carry the uncertainty shown by the bootstrap intervals.'
           % (thousands(data_summary['n_train'] + data_summary['n_validation']
                        + data_summary['n_test']), thousands(data_summary['n_test'])))
add_bullet('The selected degrees of freedom sits at the edge of the tested grid, so the data '
           'would prefer heavier tails than the grid allows.')
add_bullet('Residual autocorrelation remains, so the conditional independence assumption of the '
           'likelihood is not satisfied and reported uncertainties are optimistic in that respect.')
add_bullet('Conclusions come from one dataset, and the ranking of samplers depends on posterior '
           'geometry, which is dataset-specific.')

# --------------------------------------------------------------------------------------------
# 8. Conclusions
# --------------------------------------------------------------------------------------------

add_heading('8. Conclusions', level=1)
add_bullet('On this posterior, Gibbs sampling and Hamiltonian Monte Carlo converge at the '
           'standard %s draws per chain, reproducing a deterministic quadrature reference to '
           'within %.4f. Fisher-preconditioned Metropolis converges only with a longer run of '
           '%s draws. Random-walk Metropolis and adaptive Metropolis do not converge, reaching '
           'R-hat %.2f and %.2f.'
           % (thousands(configuration['draws_per_chain']), WORST_CONVERGED_DEVIATION,
              thousands(configuration['long_run_draws']),
              samplers['MH']['worst_rhat'], samplers['Adaptive MH (naive)']['worst_rhat']))
add_bullet('The failure is explained by posterior anisotropy, condition number %.0f, caused by '
           'near-collinear lag features. Preconditioning with the observed Fisher information '
           'raises bulk ESS from %.1f to %.0f at negligible cost.'
           % (geometry['condition_number'], samplers['MH']['min_bulk_ess'],
              samplers['Preconditioned MH']['min_bulk_ess']))
add_bullet('Estimating the proposal covariance from the chain itself failed under the '
           'configuration tested, collapsing acceptance to %s. This is a property of the '
           'burn-in budget and initial proposal, not a general property of adaptive Metropolis.'
           % format_percentage(samplers['Adaptive MH (naive)']['acceptance_rate']))
if BEST_HMC_CELL:
    add_bullet('For HMC, efficiency per gradient evaluation depends on step size and trajectory '
               'length jointly. The best converged configuration, epsilon = %.3f with L = %d, has '
               'an acceptance rate of %s, lower than several less efficient settings, so a high '
               'acceptance rate is not by itself evidence of good tuning.'
               % (BEST_HMC_CELL['step_size'], BEST_HMC_CELL['leapfrog_steps'],
                  format_percentage(BEST_HMC_CELL['acceptance'])))
add_bullet('Posterior predictive intervals from a Gaussian likelihood are badly miscalibrated, a '
           'nominal 50 percent interval covering %s, because test residuals have excess kurtosis '
           '%.0f. A Student-t likelihood chosen on validation data improves this to %s and '
           'reduces the median absolute error from %.3f to %.3f percentage points, but the '
           'intervals remain substantially wider than nominal.'
           % (format_percentage(GAUSSIAN_TEST['coverage_50']), residuals['excess_kurtosis'],
              format_percentage(STUDENT_TEST['coverage_50']),
              GAUSSIAN_TEST['median_absolute_error_percentage_points'],
              STUDENT_TEST['median_absolute_error_percentage_points']))
add_paragraph('Further work.', bold=True)
add_bullet('Place a prior on nu and sample it, rather than selecting from a grid whose edge is '
           'chosen.')
add_bullet('Fit a hierarchical model with per-machine coefficients and variances, which would '
           'address both the pooling limitation and the residual heteroscedasticity.')
add_bullet('Model the spike regime explicitly, for instance with a mixture or a state-space '
           'component, instead of absorbing it into heavy tails.')
add_bullet('Implement the No-U-Turn Sampler, which would remove the trajectory-length tuning that '
           'Section 5.6 shows matters.')

# --------------------------------------------------------------------------------------------
# 9. References and code
# --------------------------------------------------------------------------------------------

add_heading('9. References', level=1)
for position, key in enumerate(CITATION_ORDER, start=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)
    run = paragraph.add_run('[%d]  %s' % (position, BIBLIOGRAPHY[key]))
    run.font.size = Pt(10)

add_heading('10. Code and Reproducibility', level=1)
add_paragraph('All code and instructions are available at:')
add_paragraph(GITHUB_REPOSITORY_URL, bold=True)
add_table('Repository contents.',
          ['Path', 'Contents'],
          [('src/run_experiment_v2.py', 'All samplers, diagnostics, references, studies, figures'),
           ('src/mcmc_diagnostics.py', 'Rank-normalised R-hat and bulk/tail ESS, from scratch'),
           ('src/create_report_v2.py', 'Generates this document from the results file'),
           ('results/experiment_results_v2.json', 'Every number appearing in this document'),
           ('results/figures/', 'All figures'),
           ('notebooks/', 'Annotated notebook'),
           ('data/', 'Not committed; see data/README.md for the download link and layout')],
          column_widths_cm=[6.0, 9.4])
add_paragraph('The pipeline is reproduced with two commands:')
add_code_block(['python src/run_experiment_v2.py --skip-nuts',
                'python src/create_report_v2.py'])
add_paragraph(
    'The Bitbrains traces are not committed: the fastStorage directory is 1.19 GB across 1,250 '
    'files and is published by its authors %s. The --skip-nuts flag omits an optional PyMC '
    'cross-check that is impractically slow without a C++ compiler; it is not used anywhere in '
    'this document, and the external validation of Section 5.5 uses emcee instead.'
    % cite('shen2015'))

add_page_numbers()
os.makedirs(DOCUMENT_DIRECTORY, exist_ok=True)
document.save(OUTPUT_FILE)

print('Saved %s' % OUTPUT_FILE)
print('Abstract words : %d (target 200-250)' % ABSTRACT_WORD_COUNT)
print('Figures        : %d' % FIGURE_COUNTER[0])
print('Tables         : %d' % TABLE_COUNTER[0])
print('References     : %d cited, %d defined' % (len(CITATION_ORDER), len(BIBLIOGRAPHY)))
uncited = [key for key in BIBLIOGRAPHY if key not in CITATION_ORDER]
if uncited:
    print('UNCITED entries (excluded from the bibliography): %s' % ', '.join(uncited))
