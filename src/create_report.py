import sys
sys.stdout.reconfigure(encoding='utf-8')

import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES_DIR = os.path.join(BASE_DIR, 'results', 'figures')
DOCS_ARCHIVE_DIR = os.path.join(BASE_DIR, 'docs', 'archive')
os.makedirs(DOCS_ARCHIVE_DIR, exist_ok=True)

with open(os.path.join(BASE_DIR, 'results', 'experiment_results.json'), 'r') as f:
    results = json.load(f)

mh = results['MH']
gibbs = results['Gibbs']
hmc = results['HMC']
ols_rmse = results['ols_rmse']
rhat = results['rhat']
mh_sens = results['mh_sensitivity']
hmc_sens = results['hmc_sensitivity']

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles['Heading %d' % level]
    heading_style.font.color.rgb = RGBColor(0x00, 0x52, 0x8A)

FIGURE_NUMBER = [0]


def add_paragraph(text, bold=False, italic=False, size=11, alignment=None, space_after=6):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p


def add_figure(filename, caption, width_inches=5.5):
    filepath = os.path.join(FIGURES_DIR, filename)
    if os.path.exists(filepath):
        FIGURE_NUMBER[0] += 1
        doc.add_picture(filepath, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph(
            'Figure %d: %s' % (FIGURE_NUMBER[0], caption),
            italic=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
        )
    else:
        add_paragraph('[Figure missing: %s]' % filename, italic=True, size=10)


def add_pseudocode_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def fmt_pct(value):
    """Format a 0-1 fraction as a percentage string like '34.2%'."""
    return '%.1f%%' % (value * 100)


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()

add_paragraph(
    'MCMC Sampling Methods for Bayesian Linear Regression:\n'
    'A Comparative Study of Metropolis-Hastings, Gibbs Sampling,\n'
    'and Hamiltonian Monte Carlo',
    bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20
)

add_paragraph(
    'Predicting Cloud Server CPU Load Using the Bitbrains Datacenter Traces',
    italic=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30
)

add_paragraph(
    'Elad Dagmi & Shaked Mizrahi',
    bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
)

add_paragraph(
    'Advanced Methods in Machine Learning\nAugust 2026',
    size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30
)

doc.add_page_break()

# ============================================================
# ABSTRACT  (fix #1: Gibbs wording, fix #3: scope, fix #7: soften convergence, fix #14: wording)
# ============================================================
doc.add_heading('Abstract', level=1)

add_paragraph(
    'This project presents a comparative study of three Markov Chain Monte Carlo (MCMC) '
    'sampling methods applied to Bayesian linear regression for predicting CPU usage in '
    'cloud datacenter virtual machines. We implement Metropolis-Hastings (MH), Gibbs Sampling, '
    'and Hamiltonian Monte Carlo (HMC) from scratch using Python and NumPy, and evaluate them '
    'on the Bitbrains Datacenter Traces dataset containing telemetry from 1,250 virtual machines.'
)

add_paragraph(
    'Our comparison spans six metrics: convergence rate, acceptance rate, effective sample size '
    '(ESS), runtime efficiency (ESS/second), prediction accuracy (RMSE), and calibration '
    '(coverage of credible intervals). We run three independent chains per method with 10,000 '
    'post-burn-in samples each, and assess convergence using trace plots and the Gelman-Rubin '
    'R-hat statistic.'
)

# fix #1: no "perfect ESS" / "independent"; fix #3: CPU load focus; fix #14: academic tone
add_paragraph(
    'Results show that all three methods approximate the same posterior distribution, providing '
    'evidence for implementation correctness. Gibbs Sampling achieves 100%% acceptance and '
    'near-maximal ESS (%.0f) with an ESS/sec of %.1f. HMC attains ESS of %.1f with %s '
    'acceptance. MH is the simplest to implement but produces the lowest ESS (%.1f) due to '
    'random-walk behavior. All methods achieve similar RMSE (\u22480.29) and conservative '
    '95%% credible intervals (coverage 98.7-99.2%%). The resulting uncertainty estimates may '
    'support downstream tasks such as overload detection and risk-aware resource allocation.'
    % (gibbs['avg_ess'], gibbs['ess_per_sec'], hmc['avg_ess'],
       fmt_pct(hmc['acceptance_rate']), mh['avg_ess'])
)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)

toc_items = [
    ('1.', 'Introduction'),
    ('2.', 'Theoretical Background'),
    ('  2.1', 'Bayesian Linear Regression'),
    ('  2.2', 'Markov Chain Monte Carlo (MCMC)'),
    ('  2.3', 'Metropolis-Hastings Algorithm'),
    ('  2.4', 'Gibbs Sampling'),
    ('  2.5', 'Hamiltonian Monte Carlo'),
    ('  2.6', 'Convergence Diagnostics'),
    ('  2.7', 'Computational Complexity Analysis'),
    ('3.', 'Dataset: Bitbrains Datacenter Traces'),
    ('4.', 'Methodology'),
    ('  4.1', 'Feature Engineering'),
    ('  4.2', 'Model Specification'),
    ('  4.3', 'Sampler Implementations'),
    ('  4.4', 'Experimental Setup'),
    ('5.', 'Results'),
    ('  5.1', 'Exploratory Data Analysis'),
    ('  5.2', 'Convergence Analysis'),
    ('  5.3', 'Sampling Efficiency'),
    ('  5.4', 'Prediction Accuracy'),
    ('  5.5', 'Calibration'),
    ('  5.6', 'Hyperparameter Sensitivity'),
    ('6.', 'Discussion'),
    ('7.', 'Conclusions'),
    ('8.', 'References'),
    ('9.', 'Code'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('%s  %s' % (num, title))
    run.font.size = Pt(11)
    if not num.startswith(' '):
        run.font.bold = True

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION  (fix #3: scope, fix #10: MCMC motivation, fix #14: wording)
# ============================================================
doc.add_heading('1. Introduction', level=1)

# fix #10: explicit connection between Bayesian inference and MCMC
add_paragraph(
    'In Bayesian machine learning, inference requires computing posterior distributions over '
    'model parameters given observed data. For most practical models, these posterior distributions '
    'are analytically intractable, as the normalization integral cannot be computed in closed form. '
    'Markov Chain Monte Carlo (MCMC) methods address this challenge by constructing a Markov chain '
    'whose stationary distribution equals the desired posterior, enabling approximate inference '
    'through sampling.'
)

add_paragraph(
    'In Bayesian linear regression, the objective is not only to estimate a single coefficient '
    'vector, but to approximate the full posterior distribution over the regression parameters. '
    'MCMC methods allow us to generate samples from this posterior, which can then be used to '
    'estimate posterior means, credible intervals, and predictive uncertainty. This makes MCMC '
    'essential whenever closed-form posteriors are unavailable or when the model lacks conjugacy.'
)

add_paragraph(
    'This project compares three prominent MCMC sampling methods: Metropolis-Hastings (MH), '
    'Gibbs Sampling, and Hamiltonian Monte Carlo (HMC). Each method represents a different '
    'approach to exploring the parameter space:'
)

add_bullet(
    'Metropolis-Hastings: The most general MCMC algorithm, using random-walk proposals '
    'with an accept/reject mechanism based on the posterior density ratio.'
)
add_bullet(
    'Gibbs Sampling: A special case of MH that samples each parameter from its full '
    'conditional distribution, achieving 100% acceptance when conjugate priors are used.'
)
add_bullet(
    'Hamiltonian Monte Carlo: A gradient-informed method that uses Hamiltonian dynamics '
    'and the leapfrog integrator to propose distant, high-probability states.'
)

# fix #3: CPU load prediction focus, not "failures"
add_paragraph(
    'We apply all three methods to Bayesian linear regression for predicting CPU usage in cloud '
    'datacenter virtual machines, using the Bitbrains Datacenter Traces dataset. This project '
    'focuses on CPU load prediction with uncertainty quantification. The resulting uncertainty '
    'estimates may support downstream tasks such as overload detection, anomaly monitoring, and '
    'risk-aware resource allocation in cloud environments.'
)

# fix #14: "demonstrates" -> "supports a deeper understanding"
add_paragraph(
    'Our goal is not only to predict CPU load, but to rigorously compare the three sampling '
    'methods across multiple dimensions: convergence speed, sampling efficiency, predictive '
    'accuracy, calibration of uncertainty estimates, and sensitivity to hyperparameters. All '
    'implementations are written from scratch (without probabilistic programming libraries) to '
    'support a deeper understanding of the algorithmic mechanisms and implementation details.'
)

# ============================================================
# 2. THEORETICAL BACKGROUND
# ============================================================
doc.add_heading('2. Theoretical Background', level=1)

doc.add_heading('2.1 Bayesian Linear Regression', level=2)

add_paragraph(
    'Bayesian linear regression models the relationship between features X and target y as:'
)
# fix #9: clean equation formatting
add_paragraph(
    'y | X, \u03b2, \u03c3\u00b2  ~  N(X\u03b2, \u03c3\u00b2I)',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)
add_paragraph(
    'where \u03b2 is the vector of regression coefficients and \u03c3\u00b2 is the noise variance. '
    'Unlike frequentist regression, which produces point estimates, the Bayesian approach places '
    'prior distributions on the parameters and computes a full posterior distribution:'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Posterior:   ')
run.font.size = Pt(11)
run.font.italic = True
run = p.add_run('p(\u03b2, \u03c3\u00b2 | y, X)')
run.font.size = Pt(11)
run.font.bold = True
run = p.add_run('  \u221d  ')
run.font.size = Pt(11)
run = p.add_run('p(y | X, \u03b2, \u03c3\u00b2)')
run.font.size = Pt(11)
run.font.bold = True
run = p.add_run('  \u00d7  ')
run.font.size = Pt(11)
run = p.add_run('p(\u03b2)')
run.font.size = Pt(11)
run.font.bold = True
run = p.add_run('  \u00d7  ')
run.font.size = Pt(11)
run = p.add_run('p(\u03c3\u00b2)')
run.font.size = Pt(11)
run.font.bold = True
add_paragraph('We use conjugate priors:')
add_bullet(
    '\u03b2 ~ N(0, \u03c4\u00b2I) \u2014 a zero-mean Gaussian prior with variance \u03c4\u00b2 = 10'
)
add_bullet('\u03c3\u00b2 ~ Inverse-Gamma(a\u2080, b\u2080) with a\u2080 = 2, b\u2080 = 1')
add_paragraph(
    'The conjugate structure ensures that the full conditional distributions (needed for Gibbs '
    'Sampling) have closed-form expressions, while also providing a differentiable log-posterior '
    '(needed for HMC).'
)

doc.add_heading('2.2 Markov Chain Monte Carlo (MCMC)', level=2)

add_paragraph(
    'MCMC methods generate a sequence of samples {\u03b8\u2081, \u03b8\u2082, ..., \u03b8_T} from a '
    'Markov chain whose stationary distribution is the target posterior \u03c0(\u03b8). The key '
    'theoretical guarantee is that if the chain satisfies detailed balance:'
)
add_paragraph(
    '\u03c0(x) \u00b7 T(x \u2192 y)  =  \u03c0(y) \u00b7 T(y \u2192 x)',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)
add_paragraph(
    'then the chain is reversible with respect to \u03c0, and under mild conditions (ergodicity, '
    'aperiodicity), the chain converges to \u03c0 regardless of the initial state. The initial '
    'samples before convergence are discarded as "burn-in."'
)

doc.add_heading('2.3 Metropolis-Hastings Algorithm', level=2)

add_paragraph(
    'The Metropolis-Hastings algorithm proposes a new state \u03b8\' from a proposal distribution '
    'q(\u03b8\'|\u03b8) and accepts it with probability:'
)
add_paragraph(
    '\u03b1 = min(1,  [\u03c0(\u03b8\') \u00b7 q(\u03b8|\u03b8\')] / [\u03c0(\u03b8) \u00b7 q(\u03b8\'|\u03b8)])',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)
add_paragraph(
    'For a symmetric random-walk proposal (q(\u03b8\'|\u03b8) = q(\u03b8|\u03b8\')), this simplifies to '
    '\u03b1 = min(1, \u03c0(\u03b8\')/\u03c0(\u03b8)). The algorithm is simple but suffers from random-walk '
    'behavior: consecutive samples are highly correlated, and the step size must be carefully tuned. '
    'Too large a step size leads to frequent rejections; too small leads to slow exploration. The '
    'optimal acceptance rate for a d-dimensional Gaussian target is approximately 23.4% '
    '(Roberts et al., 1997).'
)
add_paragraph(
    'In our implementation, we sample \u03c3\u00b2 on the log scale (log(\u03c3\u00b2)) to ensure '
    'positivity, applying the appropriate Jacobian correction to the acceptance ratio.'
)

add_paragraph('Algorithm 1: Metropolis-Hastings', bold=True, size=11, space_after=2)
add_pseudocode_block([
    'Input: initial \u03b8\u2080, step sizes \u03c3_\u03b2, \u03c3_\u03c3, iterations T, burn-in B',
    'for t = 1 to T + B do:',
    '    // Propose new state',
    '    \u03b2\' \u2190 \u03b2 + \u03b5_\u03b2,  where \u03b5_\u03b2 ~ N(0, \u03c3_\u03b2\u00b2 I)',
    '    log(\u03c3\u00b2)\' \u2190 log(\u03c3\u00b2) + \u03b5_\u03c3,  where \u03b5_\u03c3 ~ N(0, \u03c3_\u03c3\u00b2)',
    '    \u03c3\u00b2\' \u2190 exp(log(\u03c3\u00b2)\')',
    '    // Compute log acceptance ratio (with Jacobian)',
    '    log \u03b1 \u2190 log p(y|X,\u03b2\',\u03c3\u00b2\') + log p(\u03b2\') + log p(\u03c3\u00b2\')',
    '          - log p(y|X,\u03b2,\u03c3\u00b2) - log p(\u03b2) - log p(\u03c3\u00b2)',
    '          + log(\u03c3\u00b2\') - log(\u03c3\u00b2)       // Jacobian correction',
    '    u ~ Uniform(0, 1)',
    '    if log(u) < log \u03b1 then:',
    '        \u03b8 \u2190 \u03b8\'     // Accept',
    '    // else: keep current \u03b8 (reject)',
    'return {\u03b8_B+1, ..., \u03b8_T+B}',
])
add_paragraph('')

doc.add_heading('2.4 Gibbs Sampling', level=2)

add_paragraph(
    'Gibbs Sampling is a special case of MH where each parameter is sampled from its full '
    'conditional distribution. For Bayesian linear regression with conjugate priors, the full '
    'conditionals have closed-form expressions:'
)
add_paragraph(
    '\u03b2 | \u03c3\u00b2, y, X  ~  N(\u03bc_\u03b2, \u03a3_\u03b2)',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
)
add_paragraph(
    'where  \u03a3_\u03b2 = (X\u1d40X / \u03c3\u00b2 + I / \u03c4\u00b2)\u207b\u00b9  and  '
    '\u03bc_\u03b2 = \u03a3_\u03b2 \u00b7 (X\u1d40y / \u03c3\u00b2)',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)
add_paragraph(
    '\u03c3\u00b2 | \u03b2, y, X  ~  Inverse-Gamma(a_n, b_n)',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
)
add_paragraph(
    'where  a_n = a\u2080 + n/2  and  b_n = b\u2080 + ||y - X\u03b2||\u00b2 / 2',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)
# fix #1, #11: accurate description of Gibbs acceptance
add_paragraph(
    'Because we sample from the exact conditionals, every proposal is accepted (100% acceptance '
    'rate). This 100% acceptance is a structural property of the algorithm when conjugate priors '
    'are available, and should not be interpreted as direct evidence of superior sampling quality '
    'compared to MH or HMC. Gibbs can suffer from slow mixing when parameters are strongly '
    'correlated, since each parameter is updated one at a time while holding the others fixed.'
)

add_paragraph('Algorithm 2: Gibbs Sampling', bold=True, size=11, space_after=2)
add_pseudocode_block([
    'Input: initial \u03b2\u2080, \u03c3\u00b2\u2080, iterations T, burn-in B',
    'Pre-compute: X\u1d40X, X\u1d40y',
    'for t = 1 to T + B do:',
    '    // Sample \u03b2 from full conditional',
    '    \u03a3_\u03b2 \u2190 (X\u1d40X / \u03c3\u00b2 + I / \u03c4\u00b2)\u207b\u00b9',
    '    \u03bc_\u03b2 \u2190 \u03a3_\u03b2 \u00b7 (X\u1d40y / \u03c3\u00b2)',
    '    \u03b2 ~ N(\u03bc_\u03b2, \u03a3_\u03b2)',
    '    // Sample \u03c3\u00b2 from full conditional',
    '    r \u2190 y - X\u03b2',
    '    a_n \u2190 a\u2080 + n/2',
    '    b_n \u2190 b\u2080 + r\u1d40r / 2',
    '    \u03c3\u00b2 ~ Inverse-Gamma(a_n, b_n)',
    'return {\u03b8_B+1, ..., \u03b8_T+B}',
])
add_paragraph('')

doc.add_heading('2.5 Hamiltonian Monte Carlo', level=2)

add_paragraph(
    'Hamiltonian Monte Carlo (HMC) augments the parameter space with auxiliary momentum variables '
    'and simulates Hamiltonian dynamics to make proposals that can move far from the current state '
    'while maintaining a high acceptance probability. The Hamiltonian is:'
)
add_paragraph(
    'H(q, p) = U(q) + K(p) = -log \u03c0(q) + p\u1d40p / 2',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)
add_paragraph(
    'where q represents the parameters, p the momentum, U(q) the potential energy (negative '
    'log-posterior), and K(p) the kinetic energy. The leapfrog integrator simulates the dynamics '
    'in L discrete steps of size \u03b5:'
)
add_paragraph('Leapfrog Integration (repeated L times):', bold=True)
add_bullet('Step 1 \u2014 Half-step momentum:   p  \u2190  p  \u2212  (\u03b5 / 2) \u00b7 \u2207U(q)')
add_bullet('Step 2 \u2014 Full-step position:    q  \u2190  q  +  \u03b5 \u00b7 p')
add_bullet('Step 3 \u2014 Half-step momentum:   p  \u2190  p  \u2212  (\u03b5 / 2) \u00b7 \u2207U(q)')
add_bullet('After L leapfrog steps, apply Metropolis accept/reject using \u0394H')

add_paragraph(
    'The leapfrog integrator is symplectic, meaning it preserves volume in phase space and is '
    'time-reversible. These properties ensure that the acceptance probability remains high even '
    'for large moves. The gradient \u2207U(q) = -\u2207 log \u03c0(q) guides proposals toward '
    'high-probability regions, dramatically reducing random-walk behavior compared to MH.'
)
add_paragraph(
    'HMC has two hyperparameters: step size \u03b5 and number of leapfrog steps L. The optimal '
    'acceptance rate is approximately 65-80% (Beskos et al., 2013). In our implementation, we '
    'parameterize \u03c3\u00b2 as exp(log \u03c3\u00b2) for unconstrained sampling and compute '
    'analytical gradients of the log-posterior.'
)

add_paragraph('Algorithm 3: Hamiltonian Monte Carlo', bold=True, size=11, space_after=2)
add_pseudocode_block([
    'Input: initial q\u2080, step size \u03b5, leapfrog steps L, iterations T, burn-in B',
    'for t = 1 to T + B do:',
    '    p ~ N(0, I)                    // Resample momentum',
    '    q\', p\' \u2190 q, p',
    '    // Leapfrog integration',
    '    for l = 1 to L do:',
    '        p\' \u2190 p\' - (\u03b5/2) \u2207U(q\')     // Half-step momentum',
    '        q\' \u2190 q\' + \u03b5 \u00b7 p\'              // Full-step position',
    '        p\' \u2190 p\' - (\u03b5/2) \u2207U(q\')     // Half-step momentum',
    '    // Metropolis accept/reject',
    '    \u0394H \u2190 H(q\', p\') - H(q, p)',
    '    u ~ Uniform(0, 1)',
    '    if log(u) < -\u0394H then:',
    '        q \u2190 q\'                      // Accept',
    'return {q_B+1, ..., q_T+B}',
])
add_paragraph('')

doc.add_heading('2.6 Convergence Diagnostics', level=2)

add_paragraph('We use several diagnostics to assess convergence and sampling quality:')

add_paragraph('Trace Plots:', bold=True)
add_paragraph(
    'Visual inspection of the sampled parameter values across iterations. A well-mixed chain '
    'shows rapid oscillation around a stable mean without trends or long-range patterns.'
)
add_paragraph('Gelman-Rubin R-hat Statistic:', bold=True)
# fix #9: clean R-hat formula
add_paragraph(
    'Compares between-chain and within-chain variance across M independent chains. The statistic '
    'is defined as R-hat = sqrt(V-hat / W), where V-hat is the pooled variance estimate and W '
    'is the within-chain variance. Values close to 1.0 indicate convergence; R-hat < 1.1 is the '
    'standard threshold.'
)
add_paragraph('Effective Sample Size (ESS):', bold=True)
add_paragraph(
    'Accounts for autocorrelation in the chain. ESS = n / (1 + 2 \u2211 \u03c1_k), where \u03c1_k '
    'is the autocorrelation at lag k. Higher ESS means more independent information per sample. '
    'ESS/second combines sampling quality with computational cost.'
)

doc.add_heading('2.7 Computational Complexity Analysis', level=2)

add_paragraph(
    'Understanding the per-iteration cost of each sampler is essential for interpreting runtime '
    'results. Let n denote the number of data points, p the number of parameters (features + 1), '
    'and L the number of leapfrog steps in HMC.'
)

add_paragraph('Time Complexity:', bold=True, size=12)

add_paragraph('Metropolis-Hastings: O(np) per iteration', bold=True)
add_paragraph(
    'Each MH iteration requires computing the log-likelihood, which involves the matrix-vector '
    'product X\u03b2 in O(np) and the residual sum of squares in O(n). The proposal and acceptance '
    'check are O(p). The dominant cost is the likelihood evaluation, giving O(np) per iteration.'
)
add_paragraph('Gibbs Sampling: O(np + p\u00b3) per iteration', bold=True)
add_paragraph(
    'Sampling \u03b2 from the multivariate normal full conditional requires solving a p\u00d7p linear '
    'system (or computing the matrix inverse), which costs O(p\u00b3). The matrix-vector products '
    'X\u1d40X and X\u1d40y are pre-computed once in O(np\u00b2). Each iteration then costs O(p\u00b3) '
    'for the Cholesky factorization plus O(np) for the residual computation in the \u03c3\u00b2 '
    'conditional. For our problem (n=3,500, p=%d), p\u00b3 = %d \u226a np = %d, so the '
    'dominant cost is O(np).'
    % (results['n_features'], results['n_features'] ** 3,
       results['n_train'] * results['n_features'])
)
add_paragraph('Hamiltonian Monte Carlo: O(L \u00b7 np) per iteration', bold=True)
add_paragraph(
    'Each leapfrog step requires computing the gradient of the log-posterior, which involves '
    'X\u1d40(y - X\u03b2) in O(np). With L leapfrog steps per iteration, the total cost is O(Lnp). '
    'For L=15, each HMC iteration is approximately 15\u00d7 more expensive than one MH iteration. '
    'However, HMC proposals move much further in parameter space, so the effective cost per '
    'independent sample (ESS/sec) can still be favorable.'
)

add_paragraph('Space Complexity:', bold=True, size=12)

add_paragraph('Metropolis-Hastings: O(np + Tp) total', bold=True)
add_paragraph(
    'MH requires storing the data matrix X of size O(np), the current parameter vector O(p), '
    'and the chain history of T samples O(Tp). No additional large matrices are needed since '
    'the likelihood is computed directly.'
)
add_paragraph('Gibbs Sampling: O(np + p\u00b2 + Tp) total', bold=True)
add_paragraph(
    'In addition to the data matrix and chain history, Gibbs requires storing the pre-computed '
    'X\u1d40X matrix of size O(p\u00b2) and the Cholesky factor of the precision matrix O(p\u00b2). '
    'The p\u00b2 overhead is negligible when p \u226a n.'
)
add_paragraph('Hamiltonian Monte Carlo: O(np + Tp) total', bold=True)
add_paragraph(
    'HMC stores the data matrix, chain history, and additionally the momentum vector O(p) and '
    'gradient vector O(p) during leapfrog integration. The leapfrog steps are computed in-place, '
    'so no additional storage proportional to L is required. Total: O(np + Tp + p) = O(np + Tp).'
)

complexity_table = doc.add_table(rows=4, cols=5)
complexity_table.style = 'Light Grid Accent 1'
complexity_table.alignment = WD_TABLE_ALIGNMENT.CENTER
complexity_headers = ['Method', 'Time (per iter.)', 'Space (total)',
                      'Our Setting', 'Measured Time']
for i, h in enumerate(complexity_headers):
    cell = complexity_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)
complexity_data = [
    ('MH', 'O(np)', 'O(np + Tp)', 'n=3500, p=%d' % results['n_features'],
     '%.2f s' % mh['time']),
    ('Gibbs', 'O(np + p\u00b3)', 'O(np + p\u00b2 + Tp)',
     'T=10000, L=\u2014', '%.2f s' % gibbs['time']),
    ('HMC', 'O(L\u00b7np)', 'O(np + Tp)', 'L=15', '%.2f s' % hmc['time']),
]
for row_idx, (method, time_cost, space_cost, setting, measured) in enumerate(complexity_data):
    for col_idx, val in enumerate([method, time_cost, space_cost, setting, measured]):
        cell = complexity_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

add_paragraph(
    'The measured runtimes are consistent with the complexity analysis: HMC is approximately '
    '%.1f\u00d7 slower than MH and %.1f\u00d7 slower than Gibbs, reflecting the L=15 leapfrog '
    'overhead. All three methods have comparable space complexity dominated by O(np + Tp).'
    % (hmc['time'] / mh['time'], hmc['time'] / gibbs['time']),
    size=10, italic=True
)

# ============================================================
# 3. DATASET
# ============================================================
doc.add_heading('3. Dataset: Bitbrains Datacenter Traces', level=1)

add_paragraph(
    'The Bitbrains Datacenter Traces (GWA-T-12) dataset was collected from a real managed hosting '
    'datacenter operated by Bitbrains IT Services Inc. in the Netherlands. It contains performance '
    'metrics from 1,250 virtual machines running on fast SAN storage, sampled at 5-minute intervals '
    'over approximately 30 days (August-September 2013).'
)

add_paragraph('Features per VM (11 columns):', bold=True)

table = doc.add_table(rows=12, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['#', 'Feature', 'Description']
features_list = [
    ('1', 'Timestamp', 'Unix timestamp in seconds'),
    ('2', 'CPU Cores', 'Number of virtual CPU cores'),
    ('3', 'CPU Capacity (MHz)', 'Provisioned CPU capacity'),
    ('4', 'CPU Usage (MHz)', 'Actual CPU usage in MHz'),
    ('5', 'CPU Usage (%)', 'CPU utilization percentage (target variable)'),
    ('6', 'Memory Provisioned (KB)', 'Allocated memory'),
    ('7', 'Memory Usage (KB)', 'Actual memory consumption'),
    ('8', 'Disk Read (KB/s)', 'Disk read throughput'),
    ('9', 'Disk Write (KB/s)', 'Disk write throughput'),
    ('10', 'Network Received (KB/s)', 'Incoming network traffic'),
    ('11', 'Network Transmitted (KB/s)', 'Outgoing network traffic'),
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

for row_idx, (num, feat, desc) in enumerate(features_list):
    for col_idx, val in enumerate([num, feat, desc]):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

add_paragraph('')
# fix #5: clarify how 5000 samples were selected
add_paragraph(
    'We use 50 VMs for our experiments, yielding approximately 430,000 time-series data points. '
    'After feature engineering (including lag and rolling features computed from past values only), '
    'observations were sorted chronologically across all VMs. The earliest 5,000 observations were '
    'selected to preserve temporal order, and the data was split into %d training and %d test '
    'samples without shuffling, ensuring that the test set follows the training set in time.'
    % (results['n_train'], results['n_test'])
)

add_figure('eda_histograms.png', 'Distribution of key features in the Bitbrains dataset.')

add_figure('eda_timeseries.png',
           'Time series of CPU usage showing temporal patterns and variability.')

add_figure('eda_correlation.png',
           'Correlation matrix of all original telemetry variables before feature selection. '
           'CPU Usage (MHz) was subsequently excluded from the model to avoid data leakage '
           'with the target variable CPU Usage (%).')

add_paragraph(
    'Citation: S. Shen, V. van Beek, and A. Iosup, "Statistical Characterization of '
    'Business-Critical Workloads Hosted in Cloud Datacenters," CCGrid 2015.',
    italic=True, size=10
)

# ============================================================
# 4. METHODOLOGY
# ============================================================
doc.add_heading('4. Methodology', level=1)

doc.add_heading('4.1 Feature Engineering', level=2)

add_paragraph(
    'Starting from the raw telemetry features, we engineer the following predictors for CPU usage:'
)
# fix #2: no CPU_Usage_MHz, explain leakage avoidance
add_bullet(
    'Raw features: Memory Usage, Disk Read/Write throughput, Network Recv/Trans throughput'
)
add_bullet('Lag features: CPU usage at t-1, t-2, t-3 (previous 5, 10, 15 minutes)')
add_bullet('Rolling statistics: 30-minute rolling mean and standard deviation of CPU usage')
add_bullet(
    'Total: 10 predictor features + intercept = %d parameters' % results['n_features']
)
add_paragraph(
    'To avoid data leakage, all lag and rolling features were computed using past observations '
    'only. CPU Usage (MHz) was excluded from the predictor set because it directly encodes the '
    'same quantity as the target variable CPU Usage (%) at the same timestamp (MHz / Capacity = '
    '%). Only features available before the prediction time were used.'
)
add_paragraph(
    'All features are standardized using z-score normalization (zero mean, unit variance) to '
    'improve numerical stability of the MCMC samplers and make the prior scale-invariant. '
    'The target variable is also standardized.'
)

doc.add_heading('4.2 Model Specification', level=2)

add_paragraph('We use Bayesian linear regression with the following specification:')
add_bullet('Likelihood: y | X, \u03b2, \u03c3\u00b2 ~ N(X\u03b2, \u03c3\u00b2I)')
add_bullet('Prior on coefficients: \u03b2 ~ N(0, 10I) \u2014 weakly informative')
add_bullet('Prior on variance: \u03c3\u00b2 ~ Inverse-Gamma(2, 1) \u2014 weakly informative')
add_paragraph(
    'The conjugate prior structure (Normal-Inverse-Gamma) was chosen deliberately to enable fair '
    'comparison: it allows Gibbs Sampling to use exact full conditionals while also being '
    'differentiable for HMC. The weak priors ensure that the data dominates the posterior.'
)

# fix #6: mention baseline
add_paragraph(
    'As a reference point, we include a standard Ordinary Least Squares (OLS) regression as a '
    'frequentist baseline (RMSE = %.4f on the test set). The Bayesian methods achieve comparable '
    'RMSE, confirming that the main advantage of the Bayesian approach is not superior point '
    'prediction accuracy, but uncertainty quantification through posterior predictive intervals, '
    'which OLS does not provide.' % ols_rmse
)

doc.add_heading('4.3 Sampler Implementations', level=2)

# fix #14: academic wording
add_paragraph(
    'All three MCMC samplers are implemented from scratch using only NumPy and SciPy, without '
    'probabilistic programming libraries (PyMC, Stan, etc.). This supports a deeper understanding '
    'of the algorithmic mechanisms and implementation details of each method.'
)

add_paragraph('Metropolis-Hastings:', bold=True)
add_bullet(
    'Symmetric random-walk proposal: \u03b8\' = \u03b8 + \u03b5, \u03b5 ~ N(0, \u03c3_prop\u00b2)'
)
add_bullet('\u03c3\u00b2 sampled on log scale with Jacobian correction')
add_bullet('Step sizes: step_beta=0.001, step_sigma2=0.05 (tuned for ~44% acceptance)')

add_paragraph('Gibbs Sampling:', bold=True)
add_bullet(
    'Alternates between sampling \u03b2|\u03c3\u00b2 and \u03c3\u00b2|\u03b2'
)
add_bullet('Uses closed-form full conditionals from conjugate structure')
add_bullet('Pre-computes X\u1d40X and X\u1d40y for efficiency')

add_paragraph('Hamiltonian Monte Carlo:', bold=True)
add_bullet('Leapfrog integrator with L=15 steps, step size \u03b5=0.002')
add_bullet(
    'Analytical gradient of log-posterior computed for \u03b2 and log(\u03c3\u00b2)'
)
add_bullet('Momentum resampled from N(0, I) at each iteration')

doc.add_heading('4.4 Experimental Setup', level=2)

add_paragraph(
    'For each method, we run %d independent chains with different random seeds. Each chain '
    'generates %s post-burn-in samples after discarding %s burn-in samples. Multiple '
    'chains enable computation of the Gelman-Rubin statistic for convergence assessment.'
    % (results['n_chains'],
       '{:,}'.format(results['n_samples']),
       '{:,}'.format(results['burn_in']))
)

# ============================================================
# 5. RESULTS
# ============================================================
doc.add_heading('5. Results', level=1)

doc.add_heading('5.1 Exploratory Data Analysis', level=2)

add_paragraph(
    'Before applying the MCMC methods, we performed exploratory data analysis on the '
    'Bitbrains dataset. Figure 1 shows the distribution of key features, revealing '
    'right-skewed distributions for CPU and memory usage, which is typical for cloud workloads '
    'where most VMs operate at low utilization with occasional spikes.'
)
add_paragraph(
    'Figure 2 displays the time series of CPU usage, showing clear temporal patterns with '
    'periodic peaks corresponding to business-hour activity cycles. Figure 3 presents the '
    'correlation matrix, revealing moderate correlations between CPU metrics and weaker '
    'correlations with memory and I/O features, justifying our multi-feature regression approach.'
)

doc.add_heading('5.2 Convergence Analysis', level=2)

add_paragraph(
    'Trace plots show the sampled values of selected parameters across iterations for all '
    'three chains per method.'
)

add_figure('trace_plots.png',
           'Trace plots of selected regression coefficients across 10,000 iterations '
           'for MH, Gibbs, and HMC. Good mixing is indicated by rapid oscillation around a '
           'stable mean.')

add_figure('trace_sigma2.png',
           'Trace plots of the noise variance parameter across iterations. Gibbs and HMC '
           'show excellent mixing, while MH exhibits visible autocorrelation.')

add_bullet(
    'Gibbs Sampling chains mix rapidly with no visible trends, indicating fast convergence.'
)
add_bullet(
    'HMC chains also show good mixing, with slightly wider exploration of the parameter space.'
)
add_bullet(
    'MH chains exhibit more autocorrelation, with visible "stickiness" where the chain '
    'remains at the same value for multiple iterations (rejected proposals).'
)

add_paragraph(
    'The Gelman-Rubin R-hat statistic confirms convergence for all methods. The table below '
    'shows R-hat values for selected parameters:'
)

rhat_table = doc.add_table(rows=4, cols=4)
rhat_table.style = 'Light Grid Accent 1'
rhat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rhat_headers = ['Parameter', 'MH', 'Gibbs', 'HMC']
for i, h in enumerate(rhat_headers):
    cell = rhat_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

rhat_params = ['Intercept', 'beta_1', 'sigma2']
rhat_labels = ['Intercept', '\u03b2\u2081', '\u03c3\u00b2']
for row_idx, (param_key, param_label) in enumerate(zip(rhat_params, rhat_labels)):
    row_data = [
        param_label,
        '%.4f' % rhat['MH'][param_key],
        '%.4f' % rhat['Gibbs'][param_key],
        '%.4f' % rhat['HMC'][param_key],
    ]
    for col_idx, val in enumerate(row_data):
        cell = rhat_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

add_paragraph('')
add_paragraph(
    'All R-hat values are below 1.03, well within the convergence threshold of 1.1. Gibbs '
    'Sampling achieves R-hat values closest to 1.000, followed by HMC and then MH. The slightly '
    'elevated R-hat for MH (Intercept: %.4f) reflects higher between-chain variability due to '
    'the random-walk behavior, but still indicates acceptable convergence.'
    % rhat['MH']['Intercept']
)

add_figure('rhat_convergence.png',
           'Running R-hat statistic over iterations for selected parameters. All methods '
           'converge below 1.1 within the first few hundred iterations.')

add_figure('posterior_distributions.png',
           'Posterior distributions of regression coefficients from all three methods. The '
           'overlapping distributions provide evidence that all samplers approximate the same '
           'posterior.')

doc.add_heading('5.3 Sampling Efficiency', level=2)

add_paragraph(
    'The table below summarizes the key efficiency metrics with actual numerical values:'
)

comp_table = doc.add_table(rows=6, cols=4)
comp_table.style = 'Light Grid Accent 1'
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
comp_headers = ['Metric', 'MH', 'Gibbs', 'HMC']
for i, h in enumerate(comp_headers):
    cell = comp_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

comp_data = [
    ('Acceptance Rate', fmt_pct(mh['acceptance_rate']),
     fmt_pct(gibbs['acceptance_rate']),
     fmt_pct(hmc['acceptance_rate'])),
    ('Avg ESS', '%.1f' % mh['avg_ess'],
     '{:,.0f}'.format(gibbs['avg_ess']),
     '{:,.1f}'.format(hmc['avg_ess'])),
    ('Runtime (3 chains)', '%.2f s' % mh['time'],
     '%.2f s' % gibbs['time'],
     '%.2f s' % hmc['time']),
    ('ESS / second', '%.1f' % mh['ess_per_sec'],
     '{:,.1f}'.format(gibbs['ess_per_sec']),
     '%.1f' % hmc['ess_per_sec']),
    ('RMSE (test set)', '%.4f' % mh['rmse'],
     '%.4f' % gibbs['rmse'],
     '%.4f' % hmc['rmse']),
]
for row_idx, (metric, mh_val, gibbs_val, hmc_val) in enumerate(comp_data):
    for col_idx, val in enumerate([metric, mh_val, gibbs_val, hmc_val]):
        cell = comp_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

add_paragraph('')

add_figure('comparison_bars.png',
           'Bar chart comparison of acceptance rate, ESS, runtime, and ESS/second '
           'across the three MCMC methods.')

add_paragraph('Key observations from the efficiency comparison:')

# fix #1: Gibbs ESS wording - not "independent", use "near-maximal"
add_bullet(
    'Gibbs Sampling achieves an ESS close to the total number of retained samples '
    '({:,}), indicating very low autocorrelation in this conjugate Bayesian linear regression '
    'setting. However, Gibbs samples are still generated sequentially by a Markov chain and are '
    'not theoretically independent. Its ESS/sec of {:,.1f} is the highest among all methods '
    'despite moderate runtime ({:.2f}s).'.format(
        int(gibbs['avg_ess']), gibbs['ess_per_sec'], gibbs['time'])
)
# fix #11: clarify acceptance rate meaning
add_bullet(
    'HMC achieves ESS of {:,.1f} with {} acceptance. While each '
    'iteration is expensive ({:.2f}s total for 3 chains), the gradient-informed proposals '
    'produce nearly uncorrelated samples, yielding ESS/sec of {:.1f}.'.format(
        hmc['avg_ess'], fmt_pct(hmc['acceptance_rate']),
        hmc['time'], hmc['ess_per_sec'])
)
add_bullet(
    'MH has the fastest runtime ({:.2f}s) but the lowest ESS ({:.1f}), resulting in '
    'ESS/sec of only {:.1f}. The {} acceptance rate is in the expected range for a '
    '{:d}-dimensional problem, but the random-walk behavior causes high autocorrelation.'.format(
        mh['time'], mh['avg_ess'], mh['ess_per_sec'],
        fmt_pct(mh['acceptance_rate']), results['n_features'])
)

# fix #11: acceptance rate interpretation paragraph
add_paragraph(
    'Acceptance rate is a meaningful tuning diagnostic for MH and HMC, as both use an explicit '
    'accept/reject step. For Gibbs Sampling, the 100% acceptance rate is a structural property '
    'of the algorithm when conjugate priors are used, and should not be interpreted as direct '
    'evidence of superior sampling quality. Therefore, Gibbs is mainly evaluated using R-hat, '
    'ESS, autocorrelation, runtime, and predictive performance.',
    italic=True, size=10
)

add_figure('autocorrelation.png',
           'Autocorrelation function across the three methods. MH shows slow '
           'decay (high autocorrelation), while Gibbs and HMC decorrelate rapidly.')

doc.add_heading('5.4 Prediction Accuracy', level=2)

add_paragraph(
    'All three methods produce nearly identical RMSE values on the test set (n=%d), which is '
    'expected since they all approximate the same posterior distribution. The posterior mean '
    'of \u03b2 (averaged across chains and post-burn-in samples) is used as the point predictor.'
    % results['n_test']
)

pred_table = doc.add_table(rows=5, cols=3)
pred_table.style = 'Light Grid Accent 1'
pred_table.alignment = WD_TABLE_ALIGNMENT.CENTER
pred_headers = ['Method', 'RMSE', '95% Coverage']
for i, h in enumerate(pred_headers):
    cell = pred_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

pred_rows = [
    ('OLS (baseline)', '%.4f' % ols_rmse, '\u2014'),
    ('MH', '%.4f' % mh['rmse'], fmt_pct(mh['coverage_95'])),
    ('Gibbs', '%.4f' % gibbs['rmse'], fmt_pct(gibbs['coverage_95'])),
    ('HMC', '%.4f' % hmc['rmse'], fmt_pct(hmc['coverage_95'])),
]
for row_idx, (method_name, rmse_val, cov_val) in enumerate(pred_rows):
    for col_idx, val in enumerate([method_name, rmse_val, cov_val]):
        cell = pred_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

add_paragraph('')
# fix #7: soften "converge to correct posterior"
add_paragraph(
    'The near-identical RMSE values (%.4f for MH vs. %.4f for Gibbs/HMC) indicate that the '
    'three samplers produce similar predictive performance. Together with the trace plots, R-hat '
    'values, and overlapping posterior distributions, this supports the conclusion that the '
    'samplers approximate the same posterior distribution. The marginal difference '
    '(\u0394RMSE = %.4f) is negligible and attributable to Monte Carlo noise from the finite '
    'sample size.' % (mh['rmse'], gibbs['rmse'], abs(mh['rmse'] - gibbs['rmse']))
)

# fix #4: Figure caption - time-series, not scatter plot
add_figure('predictions_vs_actual.png',
           'Predicted vs. actual CPU usage on the test set with 95% credible intervals. '
           'Close overlap between the predicted and actual curves indicates accurate prediction; '
           'the shaded region represents the posterior predictive uncertainty.')

doc.add_heading('5.5 Calibration', level=2)

add_paragraph(
    'Calibration measures whether the model\'s uncertainty estimates are reliable. A well-calibrated '
    'model should have approximately X% of test observations falling within the X% credible '
    'interval.'
)

cal_table = doc.add_table(rows=4, cols=3)
cal_table.style = 'Light Grid Accent 1'
cal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cal_headers = ['Method', '95% CI Coverage', '50% CI Coverage']
for i, h in enumerate(cal_headers):
    cell = cal_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

for row_idx, (method_name, method_data) in enumerate(
        [('MH', mh), ('Gibbs', gibbs), ('HMC', hmc)]):
    row_data = [method_name,
                fmt_pct(method_data['coverage_95']),
                fmt_pct(method_data['coverage_50'])]
    for col_idx, val in enumerate(row_data):
        cell = cal_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

add_paragraph('')
# fix #13: sharper calibration interpretation
cov95_min = min(mh['coverage_95'], gibbs['coverage_95'], hmc['coverage_95']) * 100
cov95_max = max(mh['coverage_95'], gibbs['coverage_95'], hmc['coverage_95']) * 100
cov50_min = min(mh['coverage_50'], gibbs['coverage_50'], hmc['coverage_50']) * 100
cov50_max = max(mh['coverage_50'], gibbs['coverage_50'], hmc['coverage_50']) * 100
add_paragraph(
    'The results indicate conservative uncertainty estimates rather than well-calibrated ones. '
    'The 95%% credible interval coverage (%.1f%%-%.1f%%) exceeds the nominal level, and the '
    '50%% interval coverage (%.1f%%-%.1f%%) is substantially above 50%%, indicating that the '
    'posterior predictive intervals are wider than necessary. This conservatism is typical for '
    'Bayesian linear regression with weakly informative priors, as the model accounts for '
    'parameter uncertainty in addition to noise variance. While over-coverage is acceptable '
    'for risk-aware applications (where underestimating uncertainty is more dangerous than '
    'overestimating it), the intervals could be tightened by using more informative priors or '
    'a larger training sample.'
    % (cov95_min, cov95_max, cov50_min, cov50_max)
)

doc.add_heading('5.6 Hyperparameter Sensitivity', level=2)

add_paragraph(
    'We analyze the sensitivity of MH and HMC to their respective hyperparameters. '
    'Gibbs Sampling has no tuning parameters when conjugate priors are used.'
)

add_paragraph('MH Proposal Step Size Sensitivity:', bold=True)

mh_sens_table = doc.add_table(rows=len(mh_sens) + 1, cols=2)
mh_sens_table.style = 'Light Grid Accent 1'
mh_sens_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Step Size (\u03c3_\u03b2)', 'Acceptance Rate']):
    cell = mh_sens_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)
for row_idx, entry in enumerate(mh_sens):
    row_data = ['%.3f' % entry['step'], fmt_pct(entry['acceptance'])]
    for col_idx, val in enumerate(row_data):
        cell = mh_sens_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

add_paragraph('')
add_paragraph(
    'The sensitivity analysis clearly shows the step-size trade-off: at \u03c3_\u03b2=0.001, '
    'the acceptance rate is %s, providing reasonable exploration. Increasing the step size '
    'to 0.01 drops acceptance to %s, and at 0.05 it falls to %s. For our '
    '%d-dimensional problem, the optimal step size is approximately 0.001.'
    % (fmt_pct(mh_sens[0]['acceptance']), fmt_pct(mh_sens[2]['acceptance']),
       fmt_pct(mh_sens[4]['acceptance']), results['n_features'])
)

add_paragraph('HMC Leapfrog Steps Sensitivity:', bold=True)

hmc_sens_table = doc.add_table(rows=len(hmc_sens) + 1, cols=2)
hmc_sens_table.style = 'Light Grid Accent 1'
hmc_sens_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Leapfrog Steps (L)', 'Acceptance Rate']):
    cell = hmc_sens_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)
for row_idx, entry in enumerate(hmc_sens):
    row_data = ['%d' % entry['leapfrog'], fmt_pct(entry['acceptance'])]
    for col_idx, val in enumerate(row_data):
        cell = hmc_sens_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

add_paragraph('')
add_paragraph(
    'HMC acceptance rates remain high and stable across leapfrog steps (94.9-97.6%%), with '
    'L=20 achieving the highest acceptance (%s). The consistently high acceptance rate '
    'indicates stable numerical integration, though it may also suggest that the step size '
    '\u03b5=0.002 is somewhat conservative; a larger step size could reduce the number of '
    'leapfrog steps needed per proposal at the cost of slightly lower acceptance. More '
    'leapfrog steps allow the proposal to travel further in parameter space at the cost of '
    'additional gradient evaluations.'
    % fmt_pct(hmc_sens[3]['acceptance'])
)

add_figure('sensitivity.png',
           'Sensitivity analysis: MH acceptance rate vs. proposal step size (left) and '
           'HMC acceptance rate vs. number of leapfrog steps (right).')

# ============================================================
# 6. DISCUSSION  (fix #1, #11, #12, #14)
# ============================================================
doc.add_heading('6. Discussion', level=1)

add_paragraph('Advantages and Limitations of Each Method:', bold=True, size=12)

add_paragraph('Metropolis-Hastings:', bold=True)
add_bullet(
    'Advantages: Simple to implement, no gradient required, works with any target density'
)
add_bullet(
    'Limitations: Random-walk behavior leads to high autocorrelation and low ESS (%.1f); '
    'step size requires careful tuning; scales poorly with dimensionality'
    % mh['avg_ess']
)

add_paragraph('Gibbs Sampling:', bold=True)
add_bullet(
    'Advantages: 100%% acceptance rate with conjugate priors; no tuning parameters; '
    'achieves the best efficiency in this conjugate model setting ({:,.1f} ESS/sec)'.format(
        gibbs['ess_per_sec'])
)
add_bullet(
    'Limitations: Requires known full conditional distributions (limits applicability); '
    'can be slow with strong inter-parameter correlations; updates one variable at a time'
)

add_paragraph('Hamiltonian Monte Carlo:', bold=True)
add_bullet(
    'Advantages: Gradient-informed proposals suppress random walk; high ESS ({:,.1f}); '
    'scales better to high dimensions; explores the posterior efficiently'.format(
        hmc['avg_ess'])
)
add_bullet(
    'Limitations: Requires differentiable log-posterior; each iteration is expensive '
    '(gradient + L leapfrog steps); two hyperparameters to tune (\u03b5 and L)'
)

add_paragraph('Efficiency Comparison Summary:', bold=True, size=12)
# fix #14: "dominates" -> "achieves the best efficiency"
add_paragraph(
    'When considering ESS per second as the primary efficiency metric, Gibbs Sampling achieves '
    'the best efficiency with {:,.1f} ESS/sec, followed by HMC at {:.1f} ESS/sec, and MH at '
    '{:.1f} ESS/sec. This ranking reflects the fundamental algorithmic differences: Gibbs '
    'exploits conjugacy for near-independent samples, HMC uses gradients for efficient '
    'exploration, and MH relies on uninformed random-walk proposals.'.format(
        gibbs['ess_per_sec'], hmc['ess_per_sec'], mh['ess_per_sec'])
)

add_paragraph('Practical Recommendations:', bold=True, size=12)
add_paragraph(
    'For models with conjugate priors, Gibbs Sampling is the natural choice due to its simplicity '
    'and guaranteed acceptance. For higher-dimensional or non-conjugate models, HMC is preferred '
    'despite its computational overhead, as the ESS per second often exceeds MH. The NUTS '
    '(No-U-Turn Sampler) extension of HMC eliminates the need to tune L and is implemented in '
    'modern probabilistic programming libraries like PyMC and Stan.'
)

# fix #12: strengthened limitations
add_paragraph('Limitations of This Study:', bold=True, size=12)
add_bullet(
    'The use of a linear model limits the ability to capture non-linear workload patterns; '
    'results may differ for more complex models such as Bayesian neural networks.'
)
add_bullet(
    'The experiment uses a subset of the Bitbrains dataset (5,000 observations from 50 VMs) '
    'for computational tractability, so the results may not fully represent large-scale '
    'deployment conditions.'
)
add_bullet(
    'Since the earliest 5,000 observations were selected to preserve temporal order, the '
    'subset may not capture all workload patterns across the full 30-day trace period, such '
    'as weekly cycles or end-of-month processing peaks.'
)
add_bullet(
    'The comparison is on a single dataset; generalization requires testing on diverse workloads.'
)
add_bullet(
    'We did not implement adaptive variants such as NUTS (for HMC) or adaptive Metropolis '
    'methods; hyperparameter tuning was manual rather than automatic.'
)

# ============================================================
# 7. CONCLUSIONS
# ============================================================
doc.add_heading('7. Conclusions', level=1)

add_paragraph(
    'This project presents a rigorous comparison of three MCMC sampling methods for Bayesian '
    'linear regression applied to cloud server CPU load prediction. Our key findings are:'
)

# fix #7, #14: softer convergence claims
add_bullet(
    'All three methods approximate the same posterior distribution (R-hat < 1.03 for all '
    'parameters), providing evidence for implementation correctness.'
)
# fix #1, #14: Gibbs wording
add_bullet(
    'Gibbs Sampling achieves 100%% acceptance and the best efficiency in this conjugate '
    'setting ({:,.1f} ESS/sec), with near-maximal ESS.'.format(gibbs['ess_per_sec'])
)
add_bullet(
    'HMC produces high ESS ({:,.1f}) with {} acceptance due to gradient-informed '
    'proposals, making it well-suited for non-conjugate or high-dimensional problems.'.format(
        hmc['avg_ess'], fmt_pct(hmc['acceptance_rate']))
)
add_bullet(
    'MH is the simplest to implement and the most general, but suffers from random-walk '
    'behavior that reduces sampling efficiency (ESS/sec = %.1f).' % mh['ess_per_sec']
)
add_bullet(
    'All methods produce conservative uncertainty estimates (95%% coverage: %.1f-%.1f%%), '
    'demonstrating the value of Bayesian inference for risk-aware applications where '
    'underestimating uncertainty is more costly than overestimating it.'
    % (cov95_min, cov95_max)
)
add_bullet(
    'Computational complexity analysis confirms that MH has the lowest per-iteration cost O(np), '
    'Gibbs costs O(np + p\u00b3), and HMC costs O(Lnp). The measured runtimes (%.2fs, %.2fs, '
    '%.2fs respectively) are consistent with this analysis.'
    % (mh['time'], gibbs['time'], hmc['time'])
)

add_paragraph(
    'Future work could extend this comparison to non-linear models (e.g., Bayesian neural '
    'networks), implement adaptive methods (NUTS for HMC, adaptive MH), apply the framework '
    'to larger-scale datasets, and explore the effect of different prior specifications on '
    'posterior inference and calibration.'
)

# ============================================================
# 8. REFERENCES
# ============================================================
doc.add_heading('8. References', level=1)

references = [
    'Bishop, C. M. (2006). Pattern Recognition and Machine Learning. Springer.',
    'Brooks, S., Gelman, A., Jones, G., & Meng, X. L. (2011). Handbook of Markov Chain '
    'Monte Carlo. CRC Press.',
    'Gelman, A., Carlin, J. B., Stern, H. S., Dunson, D. B., Vehtari, A., & Rubin, D. B. '
    '(2013). Bayesian Data Analysis (3rd ed.). CRC Press.',
    'Hastings, W. K. (1970). Monte Carlo sampling methods using Markov chains and their '
    'applications. Biometrika, 57(1), 97-109.',
    'Metropolis, N., Rosenbluth, A. W., Rosenbluth, M. N., Teller, A. H., & Teller, E. '
    '(1953). Equation of state calculations by fast computing machines. The Journal of '
    'Chemical Physics, 21(6), 1087-1092.',
    'Neal, R. M. (2011). MCMC using Hamiltonian dynamics. In Handbook of Markov Chain '
    'Monte Carlo (pp. 113-162). CRC Press.',
    'Roberts, G. O., Gelman, A., & Gilks, W. R. (1997). Weak convergence and optimal '
    'scaling of random walk Metropolis algorithms. The Annals of Applied Probability, '
    '7(1), 110-120.',
    'Shen, S., van Beek, V., & Iosup, A. (2015). Statistical Characterization of '
    'Business-Critical Workloads Hosted in Cloud Datacenters. In Proceedings of CCGrid 2015.',
    'Beskos, A., Pillai, N., Roberts, G., Sanz-Serna, J. M., & Stuart, A. (2013). '
    'Optimal tuning of the hybrid Monte Carlo algorithm. Bernoulli, 19(5A), 1501-1534.',
]

for i, ref in enumerate(references):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run('[%d]  %s' % (i + 1, ref))
    run.font.size = Pt(10)

# ============================================================
# 9. CODE
# ============================================================
doc.add_heading('9. Code', level=1)

add_paragraph(
    'The complete source code for this project is provided in the accompanying Jupyter notebook:'
)
add_paragraph(
    'Sampling_Project.ipynb',
    bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
)
add_paragraph(
    'The notebook contains all implementations from scratch using Python, NumPy, and SciPy, '
    'organized into the following sections:'
)
add_bullet('Section 1: Imports and setup')
add_bullet('Section 2: Data loading and exploratory data analysis')
add_bullet('Section 3: Feature engineering and preprocessing')
add_bullet('Section 4: Bayesian linear regression model definition (log-posterior, gradient)')
add_bullet('Section 5: Metropolis-Hastings sampler implementation')
add_bullet('Section 6: Gibbs Sampling implementation')
add_bullet('Section 7: Hamiltonian Monte Carlo implementation')
add_bullet('Section 8: Running all samplers (3 chains x 10,000 samples each)')
add_bullet('Section 9: Convergence diagnostics (trace plots, R-hat, autocorrelation)')
add_bullet('Section 10: Results comparison and prediction evaluation')
add_bullet('Section 11: Conclusions and summary tables')

add_paragraph(
    'Additionally, run_experiment.py provides a standalone script to reproduce all experiments '
    'and generate all figures used in this report.',
    italic=True, size=10
)

# ============================================================
# SAVE
# ============================================================
out = os.path.join(DOCS_ARCHIVE_DIR, 'Sampling_Project_Report_round1.docx')
doc.save(out)
print('Saved: %s' % out)
print('Note: this is the superseded first-round report. The current report is')
print('Sampling_Project_Report.docx, produced by create_report_v2.py.')
